import base64
import codecs
import datetime
import hashlib
import ipaddress
import logging
import os
import random
import re
import shutil
import socket
import sys
import tempfile
import time
import urllib
import uuid
from inspect import stack
from itertools import chain
from pathlib import Path
from shutil import copyfile
from time import sleep

import docx
import git
import requests
import xlrd
from django.apps import apps
from django.conf import settings
from django.contrib import messages
from django.core.cache import cache
from django.core.files.storage import FileSystemStorage
from django.core.mail import EmailMessage
from django.db.models.expressions import RawSQL
from django.db.utils import DataError, IntegrityError
from django.http import HttpResponseRedirect, JsonResponse
from django.shortcuts import reverse
from django.views import View
from docx import Document, opc
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor

from tests.mocks.mock_omni_check_status import mock_omni_check_status_200
from tests.mocks.mock_send_omnitracker import mock_send_omnitracker_200

try:
    from requests_ntlm import HttpNtlmAuth
except ImportError:
    logging.error("Error load requests_ntlm, run pip install requests_ntlm")

COMMIT_MESSAGE = "[ACL PORTAL] Update file:"
FUN_SPEED = 0
BASE_DIR = Path(__file__).resolve().parent.parent
LOCAL_UID = None

FORM_APPLICATION_KEYS = [
    "acl_create_info.html",
    "acl_internal_resources.html",
    "acl_dmz_resources.html",
    "acl_external_resources.html",
    "acl_traffic.html",
    "acl_approve.html",
]
FORM_URLS = [
    "acldemo_urls",
    "aclcreate_urls",
    "aclinternal_urls",
    "acldmz_urls",
    "aclexternal_urls",
    "acltraffic_urls",
    "acloverview_urls",
    "acl_approve_urls",
    "acl_pending_urls",
]
POST_FORM_KEYS = [
    "name",
    "email",
    "tel",
    "department",
    "project",
    "link",
    "d_form",
    "d_start",
    "d_complate",
    "acl_filename",
]
POST_FORM_EMPTY = ["on", "", None]
JSON_DUMPS_PARAMS = {"ensure_ascii": False}
left_rule = {"<": ":", "^": ":", ">": "-"}
right_rule = {"<": "-", "^": ":", ">": ":"}

contact_column = ["Параметр", "Значение"]
contact_table = [
    "ФИО",
    "E-mail",
    "Телефон",
    "Отдел/Управление",
    "Информационная система",
    "Описание/архитектура проекта",
    "Дата заполнения",
    "Дата ввода в эксплуатацию",
    "Дата вывода из эксплуатации",
    "Имя файла ACL в gitlab",
]
external_column = [
    "IP-адрес",
    "Полное доменное имя источника",
    "Маска подсети/Префикс",
    "Описание",
]
standart_column = ["IP-адрес", "Маска подсети/Префикс", "Описание"]
traffic_column = [
    "Hostname (Источник)",
    "IP Address (Источник)",
    "Hostname (Назначение)",
    "IP Address (Назначение)",
    "Protocol/Port (Назначение)",
    "ПО использующее нестандартный порт",
    "Описание (цель)",
    "Этот доступ резервный?",
]

OMNI_ACL_STATUS = {
    "inProgress": "JOB",
    "Done": "CMP",
    "Closed": "CMP",
    "Assigned": "JOB",
    "Registered": "JOB",
    "RejectedByUser": "CNL",
    "": "UNK",
}

OMNI_HTTP_STATUS = {
    "inProgress": "В работе",
    "Done": "Выполнено",
    "Closed": "Закрыто",
    "Assigned": "На исполнении",
    "Registered": "Новое",
    "RejectedByUser": "Отклолено",
    "UNK": "Статус не известен",
    "": "Статус не известен",
}

OMNI_HTTP_HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:85.0) Gecko/20100101 Firefox/85.0",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
    "Accept-Language": "ru-RU,ru;q=0.8,en-US;q=0.5,en;q=0.3",
    "Content-Type": "text/xml; charset=utf-8",
}


logger = logging.getLogger(__name__)


class BaseView(View):
    def dispatch(self, request, *args, **kwargs):
        if settings.DEBUG:
            return super().dispatch(request, *args, **kwargs)
        try:
            response = super().dispatch(request, *args, **kwargs)
        except Exception as e:
            logger.error(
                "{}|{}|{}|{}".format(
                    request.path,
                    str(e),
                    request.META.get("REMOTE_ADDR"),
                    datetime.datetime.today().strftime("%Y-%m-%d-%H:%M:%S"),
                )
            )
            messages.error(request, str(e))
            return HttpResponseRedirect(
                reverse("acldemo_urls")
            )  # self.__response({'errorMessage': str(e)}, status=400)

        if isinstance(response, (dict, list)):
            return self.__response(response)
        else:
            return response

    @staticmethod
    def __response(data, *, status=200):
        return JsonResponse(
            data,
            status=status,
            safe=not isinstance(data, list),
            json_dumps_params=JSON_DUMPS_PARAMS,
        )


def isvalidip(ip) -> bool:
    l = len(str(ip))
    if (l == 0) or (l > 15):
        return False
    s = str(ip).split(".")
    if len(s) >= 3:
        return True
    else:
        return False


def get_client_ip(request) -> str:
    """Получение IP адреса клиента"""
    x_forwarded_for = request.META.get("HTTP_X_FORWARDED_FOR")
    if x_forwarded_for:
        ip = x_forwarded_for.split(",")[0]
    else:
        ip = request.META.get("REMOTE_ADDR")
    return ip


def ip_status(ip=None) -> dict:
    """Проверка типа IP адресса"""
    data = {}
    data["ip"] = False
    try:
        ip = ipaddress.ip_address(ip)
    except ValueError:
        return data
    data["ip"] = True

    if ip.is_reserved:
        data["type"] = 3
        return data

    if ip.is_loopback:
        data["type"] = 4
        return data
    if ip.is_multicast:
        data["type"] = 5
        return data

    if ip.is_global:
        data["type"] = 1
    elif ip.is_private:
        data["type"] = 2
    else:
        data["type"] = 0
    return data


def request_handler(request, namespace=""):
    """Функция для заполнения глобального массива LOCAL_STORAGE из POST параметров файлов acl*"""
    INFINITY = "Нет"
    LOCAL_STORAGE = {}
    cnt_key = 0
    empty_key = 0

    if namespace == FORM_APPLICATION_KEYS[0]:  # first
        LOCAL_STORAGE[namespace] = []
        for idx, post_key in enumerate(POST_FORM_KEYS):
            if idx == len(POST_FORM_KEYS) - 1:
                if request.POST.get(post_key) in POST_FORM_EMPTY:
                    LOCAL_STORAGE[namespace].append(INFINITY)
                    continue

            if request.POST.get(post_key) not in POST_FORM_EMPTY:
                LOCAL_STORAGE[namespace].append(request.POST.get(post_key))
            else:
                LOCAL_STORAGE[namespace].append(INFINITY)

        if namespace == FORM_APPLICATION_KEYS[0]:
            if request.POST.get("action_make_docx", "") == "on":
                request.session["ACT_MAKE_DOCX"] = True
            elif "ACT_MAKE_DOCX" in request.session:
                del request.session["ACT_MAKE_DOCX"]
            if request.POST.get("action_make_git", "") == "on":
                request.session["ACT_MAKE_GIT"] = True
            elif "ACT_MAKE_GIT" in request.session:
                del request.session["ACT_MAKE_GIT"]
            if request.POST.get("action_make_omni", "") == "on":
                request.session["ACT_OMNI"] = True
            elif "ACT_OMNI" in request.session:
                del request.session["ACT_OMNI"]

            # request.session.modified = True

    else:
        if namespace == FORM_APPLICATION_KEYS[-2]:  # Traffic page
            str_pattern = "input__domain_source"
        else:
            str_pattern = "input__ip"

        for k, v in request.POST.items():
            if "input_" in str(k):
                if len(v) > 0:
                    try:
                        v = "\n".join(v.splitlines())
                        v = v.replace("\n\n", "\n")
                    except Exception as e:
                        if settings.DEBUG:
                            logger.error(f"{stack()[0][3]} {e}")
                if str_pattern in str(k):
                    if namespace in LOCAL_STORAGE:
                        LOCAL_STORAGE[namespace].append([v])
                        cnt_key += 1
                    else:
                        LOCAL_STORAGE[namespace] = [[v]]
                else:
                    if v != "":
                        LOCAL_STORAGE[namespace][cnt_key].append(v)
                    else:
                        if k == "d_complate":
                            LOCAL_STORAGE[namespace][cnt_key].append(INFINITY)
                        else:
                            empty_key += 1

        # if empty_key >= 2:
        #         del LOCAL_STORAGE[namespace]
        # return False
    return LOCAL_STORAGE


def IP2Int(ip):
    """Function convert IP to integer"""
    o = list(map(int, ip.split(".")))
    res = (16777216 * o[0]) + (65536 * o[1]) + (256 * o[2]) + o[3]
    return res


def upload_file_handler(request, functionhandler=None):
    """Функция обработки загрузки файлов и вызова функции для парсинга xls"""
    result = {}
    if "input--file--upload" in request.FILES:
        UPLOAD_PATH = tempfile.gettempdir()  # os.path.join(BASE_DIR, 'upload')
        myfile = request.FILES["input--file--upload"]
        fs = FileSystemStorage(location=UPLOAD_PATH)
        fs.save(myfile.name, myfile)
        uploaded_file_url = os.path.join(
            UPLOAD_PATH, myfile.name
        )  # bug with persone encode'
        if settings.DEBUG:
            print(f"Upload file to: {uploaded_file_url}")
    else:
        result["error"] = "Отсутствуют файлы для загрузки"
        return result
    if uploaded_file_url == "":
        return {"error": "There is error upload file"}

    if functionhandler is not None:
        result = functionhandler(uploaded_file_url)
    else:
        if "xls" not in uploaded_file_url:
            # if 'ext.' in uploaded_file_url or \
            #    'aktur' in uploaded_file_url or \
            #      'alfatrah.ru' in uploaded_file_url:
            functionhandler = ExtractDataDns
        else:
            result = ExtractDataXls(request, uploaded_file_url).execute_file_parsing()
        # race condition
    time.sleep(1)
    try:
        os.remove(uploaded_file_url)
        if settings.DEBUG:
            print(f"Удаление файла: {uploaded_file_url}")
    finally:
        uploaded_file_url = None
    if isinstance(result, int):
        if result > 0:
            return {"ok": f"Добавлено новых значений: {result}"}
    else:
        if isinstance(result, dict):
            if "LOCAL_STORAGE" in result:
                request.session["LOCAL_STORAGE"] = result.get("LOCAL_STORAGE")
                if settings.DEBUG:
                    print(result.get("META"))
                return {"status": "Шаблон загружен"}
            else:
                return result
    return {"error": "Данных для добавления - нету"}


def count_perf(f):
    """Декоратор для измерения скорости поиска"""

    def wraper(*args, **kwargs):
        global FUN_SPEED
        time_init = datetime.datetime.now()
        result = f(*args, **kwargs)
        time_end = datetime.datetime.now()
        total = time_end - time_init
        FUN_SPEED = total
        return result

    return wraper


@count_perf
def DeepSearch(request, string: str = ""):
    """Функция для анализа типа данных в запросе и поиск по структуре"""
    result, tmp = "", string
    Iplist = apps.get_model("ownerlist", "Iplist")
    Acl = apps.get_model("accesslist", "ACL")

    if re.match(r"^\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}$", tmp):
        result = Iplist.objects.filter(ipv4=tmp)
    if not result:
        # acl = Acl.objects.filter(acltext__acl_internal_resources.html__contains=)

        if re.match(r"^\d{1,3}\,\d{1,3}\,\d{1,3}\,\d{1,3}$", tmp):
            tmp = string.replace(",", ".")
        if tmp:
            result = Iplist.objects.filter(ipv4__contains=tmp)[:5]
        # else:
        #     result = Iplist.objects.filter(ipv4__contains=string)[:5]
        if result:
            messages.add_message(
                request,
                messages.INFO,
                f"По запросу {string} ничего не найдено, но мы нашли похожую информацию:",
            )
    if not result:
        if re.match(r"[a-zA-Z0-9][a-zA-Z0-9-._]{1,61}", tmp):

            result = Iplist.objects.filter(hostname=tmp)[:5]

            if not result:
                if "://" in tmp:
                    tmp = tmp.split("://")[1]

                if (
                    ".vesta.ru" in tmp
                    or ".alfastrah.ru" in tmp
                    or ".dyn.vesta.ru" in tmp
                ):
                    tmp = tmp.split(".")[0]

                result = Iplist.objects.filter(hostname__icontains=tmp)[:5]
                if result:
                    messages.add_message(
                        request,
                        messages.INFO,
                        f"По запросу {string} ничего не найдено, но мы нашли похожую информацию:",
                    )
                if not result:
                    result = Iplist.objects.filter(comment__icontains=tmp)[:5]
                    if result:
                        messages.add_message(
                            request,
                            messages.INFO,
                            f"По запросу {string} ничего не найдено, но мы нашли похожую информацию:",
                        )
                    else:
                        try:
                            tmp = socket.gethostbyname(tmp)
                        except:
                            pass
                        if tmp:
                            result = Iplist.objects.filter(ipv4__contains=tmp)[:5]
                        if result:
                            messages.add_message(
                                request,
                                messages.INFO,
                                f"По запросу {string} ничего не найдено, но мы нашли похожую информацию:",
                            )

    if re.match("[A-Za-zА-Яа-я\d\.]{3,20}", tmp):
        tmp = tmp.replace("%", " ").replace("'", " ")
        row_query = Acl.objects.filter(
            id__in=RawSQL(
                f"select id from accesslist_acl where acltext::text like '%%{tmp}%%' limit 5",
                tmp,
            )
        )
        if row_query:

            if len(row_query) > 0 and result and len(result) > 0:
                result = chain(row_query, result)
            else:
                result = row_query

    return result


def write_history(request, string, status) -> None:
    """Сохранять историю поиска, для улучшения качества поиска"""
    hc = apps.get_model("ownerlist", "HistoryCall")
    ip = apps.get_model("ownerlist", "Iplist")

    ip_object, obj = ip.objects.get_or_create(
        ipv4=request.META.get("REMOTE_ADDR")
    )  # IP-адрес пользователя
    if request.user.is_authenticated:
        User = request.user
    else:
        User = None
    hc_object = hc.objects.create(
        string=string, ipv4=ip_object, username=User, status=status
    )
    return hc_object


def search_text(request=None, string: str = "") -> dict:
    """Функция для поиска данных в БД"""
    global FUN_SPEED
    result = DeepSearch(request, string)
    context = {"SearchFor": string}
    context["Data"] = result
    context["TakeTime"] = FUN_SPEED
    context["Info"] = ""
    FUN_SPEED = 0
    write_history(request, string, bool(result))
    return context


class ExtractDataXls:
    """Основной класс для анализа xls файла"""

    def __init__(self, request=None, filename=""):
        self.ip_addr_idx = 1
        self.count_total: int = 0  # total records in db
        self.error_count: int = 0  # total errors
        self.rb = xlrd.open_workbook(
            filename, formatting_info=True, encoding_override="utf-8"
        )

        self.current_page = None
        self.sheet_tags = self.rb.sheet_names()
        self.Vlans = apps.get_model("ownerlist", "Vlans")
        self.Tags = apps.get_model("ownerlist", "Tags")
        self.Iplist = apps.get_model("ownerlist", "Iplist")
        self.Owners = apps.get_model("ownerlist", "Owners")
        self.page_headers = [
            "ответственный",
            "комменты",
            "ip address",
            "Имя сервера",
            "отвеcтвенный",
            "nat inside",
        ]
        self.fio_exclude_list = [
            "гусев",
            "оксенюк",
            "северцев",
            "егоров",
            "совинский",
            "огнивцев",
            "допиро",
            "мюлекер",
            "уволен",
            "иренов",
            "казаков",
            "куслеев",
        ]

    def execute_file_parsing(self):
        """Выбираем парсер на основе имени страницы"""
        result = 0
        for self.sheet_tag in self.sheet_tags:
            self.current_page = self.rb.sheet_by_name(self.sheet_tag)
            if self.current_page.nrows == 0:  # Count row
                if settings.DEBUG:
                    print(f"Страница {self.current_page} пустая, пропущено...")
                    return 0
            if self.sheet_tag == "VLAN DESCRIPTION":
                result += self.ExtractVlanInfo()
            elif self.sheet_tag == "VLAN_CORE_ACI":
                result += self.ExtractVlanInfo(
                    name_idx=6,
                    location_idx=3,
                    vlan_idx=2,
                    subnet_idx=4,
                    mask_idx=5,
                    tag1_idx=7,
                    tag2_idx=8,
                )
            elif self.sheet_tag == "10.255.10.0 (NGNX-Serv)":
                result += self.ExtractIPInfo(
                    domain_idx=0, ip_idx=1, owner_idx=2, comment_idx=3
                )
            elif self.sheet_tag == "213.33.175.0 _24":
                result += self.ExtractIPInfo(
                    domain_idx=1, ip_idx=2, owner_idx=4, comment_idx=5
                )
            elif self.sheet_tag == "активка 172.16.82.X":
                result += self.ExtractIPInfo(
                    domain_idx=1, ip_idx=0, owner_idx=2, comment_idx=3
                )
            elif self.sheet_tag == "195.239.64.хх":
                result += self.ExtractIPInfo(
                    domain_idx=0, ip_idx=1, owner_idx=3, comment_idx=4
                )
            elif self.current_page.ncols == 4:
                result += self.ExtractIPInfo()
            else:
                if settings.DEBUG:
                    print(
                        f"Страница содержит другое количество колонок <> 4, {self.current_page.ncols} анализируем..."
                    )
                result += self.PageStructAnalyzer(self.current_page)
        return result

    def is_row_empty(self, row) -> bool:
        """Проверяем пустая ли запись"""
        result = True
        for d in row:
            if d != "":
                result = False
                break
        return result

    def get_ip_from_page(self, page) -> str:
        """Получаем полный IP из имени страницы"""
        try:
            ip = re.findall(r"(\d{1,3})", page)
            return ".".join(ip)
        except:
            return ""

    def ExtractVlanInfo(
        self,
        name_idx=1,
        location_idx=2,
        vlan_idx=3,
        subnet_idx=4,
        mask_idx=5,
        tag1_idx=6,
        tag2_idx=7,
    ) -> int:
        """Парсер страницы с описанием VLAN"""
        row_index: int = 0
        tags: list = []
        for row_idx in range(self.current_page.nrows):
            row = self.current_page.row_values(row_idx)
            if row_idx == 0 or self.is_row_empty(row):
                continue

            if type(row[vlan_idx]) == float:
                vlan = int(round(row[vlan_idx]))
            elif type(row[vlan_idx]) == str:
                try:
                    vlan = int(round(float(row[vlan_idx])))
                except ValueError:
                    vlan = 0

            if str(row[subnet_idx]).find("/") > 0:
                subnet = str(row[subnet_idx]).split("/")
                subnet, mask = subnet[0], int(subnet[1])
            else:
                try:
                    if len(str(row[subnet_idx])) > 15:
                        subnet = str(row[subnet_idx]).split("\n")[
                            0
                        ]  # Bug fig, if a couple value in row
                    else:
                        subnet = str(row[subnet_idx])
                except ValueError:
                    subnet = 0

                try:
                    if len(str(row[mask_idx])) > 4:
                        mask = str(row[mask_idx]).split("\n")[
                            0
                        ]  # Bug fig, if a couple value in row
                        mask = int(round(float(mask)))
                    else:
                        mask = int(round(float(row[mask_idx]))) or 0
                except ValueError:
                    mask = 0

            vlan_info, _ = self.Vlans.objects.get_or_create(
                name=str(row[name_idx]),
                location=str(row[location_idx]),
                vlan=vlan,
                subnet=subnet,
                mask=mask,
            )
            # if created obj
            if _:
                self.count_total += 1

            try:
                tags.append(self.sheet_tag)
                tags.append(row[tag1_idx])
                tags.append(row[tag2_idx])

                for tag in tags:
                    if (tag != "") and len(tag) > 1:
                        if len(str(tag).split(".")) >= 3:  # If tag as Gateway's IP
                            tag = f"Gateway: {tag}"
                        tag_info, _ = self.Tags.objects.get_or_create(
                            name=str(tag).rstrip()
                        )
                        if tag_info not in vlan_info.tags.all():
                            vlan_info.tags.add(tag_info)
                            self.count_total += 1
            except:
                pass

            finally:
                tags.clear()

        if settings.DEBUG:
            print(f"Добавленно {self.count_total} новых записей в БД.")

        return self.count_total

    def ExtractIPInfo(
        self,
        domain_idx=0,
        ip_idx=1,
        owner_idx=2,
        comment_idx=3,
        stop_recurse=False,
        HasTags=[],
    ) -> int:
        tags: list = []
        self.error_count = 0
        Header_POS = 0
        created = None

        ip_addr = ""
        domain = ""
        owner = ""
        comment = ""

        self.count_total = 0
        for row_idx in range(self.current_page.nrows):
            row = self.current_page.row_values(row_idx)
            if not self.is_row_empty(row):
                # if self.sheet_tag == 'Орел':
                #     print ('')
                if len(row) >= 3:
                    if row_idx in range(0, 3):
                        if row[row_idx] in self.page_headers:  # Пропускаем заголовки
                            Header_POS = row_idx
                            continue
                if isvalidip(row[ip_idx]):
                    ip_addr = row[ip_idx]
                elif len(str(row[ip_idx])) <= 5:  # 15.0
                    try:
                        tmp = int(round(float(row[ip_idx]))) or 0
                    except ValueError:
                        tmp = 0
                        continue
                    else:
                        if tmp > 0:
                            ip_addr = self.get_ip_from_page(str(self.sheet_tag))
                            if ip_addr != "":
                                ip_addr = ip_addr + "." + str(tmp)
                else:
                    self.error_count += 1
                    if self.error_count >= 5:
                        if not stop_recurse:
                            if settings.DEBUG:
                                print(
                                    "Много ошибок на странице, провёдем анализ страницы ..."
                                )
                            self.PageStructAnalyzer(self.current_page)
                        else:
                            if settings.DEBUG:
                                print(
                                    f"********************Ошибка при анализе странице {self.sheet_tag}****************"
                                )
                        return 0
                    continue
                try:
                    if domain_idx is not None:
                        domain = row[domain_idx]
                except:
                    domain = ""

                try:
                    if owner_idx is not None:
                        owner = row[owner_idx]
                except:
                    owner = ""

                try:
                    if comment_idx is not None:
                        comment = row[comment_idx]
                except:
                    comment = ""

                if len(HasTags) > 0:
                    for tag in HasTags:
                        try:
                            tmp = row[int(tag)]
                            if tmp != "":
                                name = self.current_page.cell_value(
                                    Header_POS, int(tag)
                                )
                                if name != "":
                                    tmp = name + ":" + tmp
                                if tmp not in tags:
                                    tags.append(tmp)
                        except:
                            pass
                if self.sheet_tag not in tags:
                    tags.append(self.sheet_tag)

                # -----------------------------------------------------------------------------------------------------------------------
                try:
                    if comment and re.match(r"([а-яА-Я\.\s()]){5,}", str(comment)):
                        tmp = comment.lower().strip()
                        exists = list(filter(lambda s: s in tmp, self.fio_exclude_list))
                        # exists = any(substring in string for string in strings)
                        # if comment.lower() in self.fio_exclude_list:
                        if len(exists) > 0:
                            owner, comment = comment, owner
                except:
                    pass
                # -----------------------------------------------------------------------------------------------------------------------

                try:
                    if owner:
                        if owner.find("://") != -1:
                            owner, comment = comment, owner
                except:
                    pass

                if ip_addr != "":
                    # if not self.isvalidip(ip_addr):

                    # continue
                    # if comment_idx != -1:
                    #     print("{} {} {} {} ".format(domain, ip_addr, row[owner_idx], row[comment_idx]))
                    # else:
                    if (owner == "") or len(owner) <= 1:
                        owner_info = None
                    else:
                        # try:
                        owner_info, created = self.Owners.objects.get_or_create(
                            username=owner
                        )
                    # except:
                    # owner_info = self.Owners.get_default_owner()

                    try:
                        ip_info, created = self.Iplist.objects.get_or_create(
                            ipv4=ip_addr,
                            hostname=domain,
                            owner=owner_info,
                            comment=comment,
                        )
                    except IntegrityError:
                        ip_info = self.Iplist.objects.get(ipv4=ip_addr)
                        ip_info.ipv4 = ip_addr
                        ip_info.hostname = domain
                        ip_info.owner = owner_info
                        ip_info.comment = comment
                        ip_info.save()

                    except DataError:
                        print(
                            f"- Ошибка данных: {ip_addr} на странице: {self.sheet_tag}"
                        )
                        continue

                    if created:
                        self.count_total += 1

                    try:
                        for tag in tags:
                            if (tag != "") and len(tag) > 1:
                                tag_info, created = self.Tags.objects.get_or_create(
                                    name=str(tag).rstrip()
                                )
                                if tag_info not in ip_info.tags.all():
                                    ip_info.tags.add(tag_info)
                    except:
                        pass

                    finally:
                        # print("{} {} {} {}-> {}".format(domain, ip_addr, owner, comment, " [" + ' '.join(tags) + "]"))

                        tags.clear()
        if settings.DEBUG:
            print(f"Страница: {self.sheet_tag} записано: {self.count_total}")
        return self.count_total

    def PageStructAnalyzer(self, page, DEBUG=False) -> None:
        """Написанный на коленке анализатор данных на странице в xls"""
        is_domain = 0
        is_ip = 0
        is_owner = 0
        is_comment = 0
        is_add = 0
        col_index = {}
        # RowIndex = 0
        # RowBlackListByNumber = []
        skip_domain = False
        skip_owner = False
        skip_ip = False
        skip_commnet = False

        col_index["domain"] = None
        col_index["ip"] = None
        col_index["owner"] = None
        col_index["comment"] = None

        Tags = []
        # col_index = {}
        # if self.sheet_tag == '172.16.88.Х Avaya':
        #     print('')
        if not (page.ncols > 0):  # and page.ncols < 6
            print(
                f"Page {self.sheet_tag} has been skipped, unknown amount col {page.ncols}"
            )
        else:
            for idx_col in range(page.ncols):
                col_stat = {
                    "Unknown": 0,
                    "is_domain": 0,
                    "is_ip": 0,
                    "is_owner": 0,
                    "is_comment": 0,
                    "is_tag": 0,
                }

                # RowIndex = 0
                if idx_col >= 8:  # Skip comment in tables
                    continue

                for index, col in enumerate(page.col_values(idx_col)):
                    if col == "":
                        continue
                    if index in range(0, 2):
                        if col in self.page_headers:
                            continue
                    try:
                        if len(str(col)) <= 1:
                            continue
                        if re.match(
                            r"^\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}$", str(col)
                        ) or re.match(r"^(\d+.\d)|(\d)$", str(col)):
                            if not skip_ip:
                                col_stat["is_ip"] += 1
                                continue
                            else:
                                if idx_col not in Tags:
                                    Tags.append(idx_col)
                                continue
                    except TypeError:
                        pass

                    try:
                        if (skip_domain == False) and re.match(
                            r"([a-zA-Z0-9\-\.\\()\_]){4,}", str(col)
                        ):  # [А-я]+
                            col_stat["is_domain"] += 1
                            continue
                    except TypeError:
                        pass

                    try:
                        if (skip_owner == False) and re.match(
                            r"([а-яА-Я\.\s()]){5,}", str(col)
                        ):
                            col_stat["is_owner"] += 1
                            continue
                    except:
                        pass

                        # else:
                    if len(str(col)) > 3:
                        t = str(col).lower().find("vlan")
                        if t > 0:
                            pass
                        #     str_sum = len(col)
                        #     #col_stat['is_tag'] += 1
                        #     try:
                        #         #cut_vlan = col.split(' ')
                        #         if idx_col not in Tags:
                        #             Tags.append(idx_col)
                        #             continue
                        #     except:
                        #         pass
                        # in ['vlan', 'vlan name', 'location', ]
                        else:
                            if not skip_commnet:
                                col_stat["is_comment"] += 1
                            else:
                                if idx_col not in Tags:
                                    Tags.append(idx_col)
                                    continue

                max_v = max(col_stat, key=col_stat.get)
                if max_v == "Unknown":
                    continue
                elif max_v == "is_domain":
                    col_index["domain"] = idx_col
                    skip_domain = True
                    if idx_col in Tags:
                        Tags.remove(idx_col)
                elif max_v == "is_owner":
                    skip_owner = True
                    col_index["owner"] = idx_col
                    if idx_col in Tags:
                        Tags.remove(idx_col)
                elif max_v == "is_ip":
                    skip_ip = True
                    col_index["ip"] = idx_col
                    if idx_col in Tags:
                        Tags.remove(idx_col)
                elif not skip_commnet:
                    col_index["comment"] = idx_col
                    if idx_col in Tags:
                        Tags.remove(idx_col)

                if max_v == "is_comment":
                    if idx_col == 2:
                        if (col_stat["is_owner"]) >= 5:
                            col_index["owner"] = idx_col
                            if idx_col in Tags:
                                Tags.remove(idx_col)
                            if settings.DEBUG:
                                print(
                                    "Page: {}| Col {}| possible: [FIX] owner (d:{},i:{}, o:{},c:{},t:{})".format(
                                        self.sheet_tag,
                                        idx_col,
                                        col_stat["is_domain"],
                                        col_stat["is_ip"],
                                        col_stat["is_owner"],
                                        col_stat["is_comment"],
                                        col_stat["is_tag"],
                                    )
                                )
                            continue
                    else:
                        skip_commnet = True

                if DEBUG:
                    print(
                        "Page: {}| Col {}| possible: {}(d:{},i:{}, o:{},c:{},t:{})".format(
                            self.sheet_tag,
                            idx_col,
                            max_v,
                            col_stat["is_domain"],
                            col_stat["is_ip"],
                            col_stat["is_owner"],
                            col_stat["is_comment"],
                            col_stat["is_tag"],
                        )
                    )

        if not DEBUG:
            return self.ExtractIPInfo(
                domain_idx=col_index["domain"],
                ip_idx=col_index["ip"],
                owner_idx=col_index["owner"],
                comment_idx=col_index["comment"],
                stop_recurse=True,
                HasTags=Tags,
            )


def update_callback_status(request, taskid, name, value, event=1):
    """Функция для создания сообщения callback на запрос статуса потоками"""
    JOB = cache.get(taskid, {})
    if event == 2:
        JOB.update({name: {"done": value}})  # done
    elif event == 1:
        JOB.update({name: {"status": value}})
    else:
        JOB.update({name: {"error": value}})
        try:
            logger.error(
                "{} [error] [client {}] [user {}] {}".format(
                    str(datetime.datetime.now()),
                    request.META.get("REMOTE_ADDR"),
                    request.user,
                    value,
                )
            )
        except:
            pass
    cache.set(taskid, JOB)


def add_hyperlink(paragraph, url, text, color, underline):

    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(docx.oxml.shared.qn("r:id"), r_id)

    new_run = docx.oxml.shared.OxmlElement("w:r")
    rPr = docx.oxml.shared.OxmlElement("w:rPr")
    if color is not None:
        c = docx.oxml.shared.OxmlElement("w:color")
        c.set(docx.oxml.shared.qn("w:val"), color)
        rPr.append(c)
    if underline:
        u = docx.oxml.shared.OxmlElement("w:u")
        u.set(docx.oxml.shared.qn("w:val"), "single")
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def make_doc(
    request=None, data_set={}, fileuuid="", gitlab_repo_url="", gitlab_filename=""
) -> str:
    """Функция для генерации docx файла"""

    def normalize_whitespace(data):
        for key, value in data.items():
            if key == "acl_traffic.html":
                value = [row + [" "] if len(row) == 7 else row for row in value]
            elif key != "acl_create_info.html":
                value = [
                    [cell.strip() if cell else " " for cell in row] for row in value
                ]
            data[key] = value
        return data

    def get_gitlab_data():
        if gitlab_repo_url and gitlab_filename:
            md_content, _ = get_acl_from_gitlab(
                repo_url=gitlab_repo_url,
                file_name=gitlab_filename,
            )
            return normalize_whitespace(convert_md_to_dict(md_content)["LOCAL_STORAGE"])
        print("Не удалось получить gitlab_data_set")
        return data_set

    def prepare_file_name():
        try:
            name = data_set["acl_create_info.html"][4] or ""
            if len(name) > 10:
                name = "_".join(name.split()[:2])
            return (
                name.replace(" ", "_")
                .replace(",", "")
                .replace(".", "")
                .replace("/", "_")
            )
        except:
            return str(uuid.uuid4())

    def highlight_changes(cell, is_new, is_deleted):
        paragraphs = cell.paragraphs if hasattr(cell, "paragraphs") else [cell]

        for paragraph in paragraphs:
            for run in paragraph.runs:
                if is_new:
                    run.font.highlight_color = WD_COLOR_INDEX.GREEN
                    run.font.color.rgb = RGBColor(255, 255, 255)
                elif is_deleted:
                    run.font.strike = True
                    run.font.highlight_color = WD_COLOR_INDEX.RED
                    run.font.color.rgb = RGBColor(255, 255, 255)

    def process_contact_table(table, data, gitlab_data):
        not_found_count = 0
        for row_idx, row_data in enumerate(data):
            if row_idx >= len(table.rows) - 1:
                table.add_row()
            cell = table.cell(row_idx, 1)

            if row_idx == 5 and row_data not in ("Нет", ""):
                try:
                    for paragraph in cell.paragraphs:
                        add_hyperlink(paragraph, row_data, row_data, "0000EE", True)
                except Exception:
                    cell.text = row_data
            else:
                cell.text = row_data

            gitlab_row_data = (
                gitlab_data[row_idx] if row_idx < len(gitlab_data) else None
            )
            if str(row_data) != str(gitlab_row_data):
                highlight_changes(cell, True, False)
                if gitlab_row_data:
                    strike_paragraph = cell.add_paragraph()
                    strike_paragraph.add_run(str(gitlab_row_data))
                    highlight_changes(strike_paragraph, False, True)
                else:
                    not_found_count += 1

        if not_found_count > 0:
            logger.info(
                f"Не удалось найти {not_found_count} строк(и) страницы Контакты в gitlab"
            )

    def process_data_table(table, data_set, gitlab_data):
        def compare_rows(row1, row2):
            return all(
                str(cell1).strip() == str(cell2).strip()
                for cell1, cell2 in zip(row1, row2)  # noqa: B905
            )

        # Очистка таблицы
        for row in table.rows[1:]:
            table._tbl.remove(row._tr)

        # Обработка данных из data_set
        for row_idx, row_data in enumerate(data_set, start=1):
            table.add_row()
            for cell_idx, cell_val in enumerate(row_data):
                if cell_idx >= len(table.columns):
                    continue
                cell = table.cell(row_idx, cell_idx)
                cell.text = str(cell_val)

                # Проверяем, есть ли эта строка в gitlab_data
                if not any(
                    compare_rows(row_data, gitlab_row) for gitlab_row in gitlab_data
                ):
                    highlight_changes(cell, True, False)  # Новая строка

        # Добавление удаленных строк из gitlab_data
        for gitlab_row in gitlab_data:
            if not any(compare_rows(gitlab_row, data_row) for data_row in data_set):
                row_idx = len(table.rows)
                table.add_row()
                for cell_idx, cell_val in enumerate(gitlab_row):
                    if cell_idx >= len(table.columns):
                        continue
                    cell = table.cell(row_idx, cell_idx)
                    cell.text = str(cell_val)
                    highlight_changes(cell, False, True)  # Удаленная строка

        # Очистка пустых строк
        for row in table.rows[::-1]:
            if all(cell.text.strip() == "" for cell in row.cells):
                table._tbl.remove(row._tr)

    data_set = normalize_whitespace(data_set)
    gitlab_data_set = get_gitlab_data()

    is_in_session = request and "ACT_MAKE_DOCX" in request.session

    TEMPLATE_FILE = os.path.join(BASE_DIR, "templates//ACL.docx")
    fileuuid = prepare_file_name()

    if is_in_session:
        update_callback_status(
            request, fileuuid, "docx_download_status", "Создаем docx файл"
        )

    try:
        gitlab_filename = gitlab_filename.split(".")[:-1][0]
        APP_FILE = os.path.join(
            settings.MEDIA_ROOT,
            f"{settings.DOCX_FILE_PATH}/ACL_{fileuuid}_{gitlab_filename}.docx",
        )
    except Exception as e:
        APP_FILE = os.path.join(
            settings.MEDIA_ROOT,
            f"{settings.DOCX_FILE_PATH}/ACL_{fileuuid}.docx",
        )
        logger.error(f"Не удалось добавить gitlab_filename в название файла docx:{e}")

    doc = Document(TEMPLATE_FILE)

    if is_in_session:
        update_callback_status(
            request, fileuuid, "docx_download_status", "Записываем изменения"
        )

    doc.styles["Normal"].font.name = "Verdana"
    doc.styles["Normal"].font.size = Pt(10)

    for data_idx, data_key in enumerate(FORM_APPLICATION_KEYS):
        table = doc.tables[data_idx]
        if data_key not in data_set:
            continue

        if data_idx == 0:  # Таблица контактов
            process_contact_table(
                table, data_set[data_key], gitlab_data_set.get(data_key, [])
            )
        else:
            process_data_table(
                table, data_set[data_key], gitlab_data_set.get(data_key, [])
            )

    if is_in_session and "taskid" in request.session and request.session["taskid"]:
        p = doc.add_paragraph(request.session["taskid"])
        doc.tables[0]._element.addnext(p._p)

    if is_in_session:
        update_callback_status(
            request,
            fileuuid,
            "docx_download_status",
            f"Сохраняем файл Application_{fileuuid}",
        )
    doc.save(APP_FILE)
    return APP_FILE


def is_valid_uuid(uuid_to_test, version=4):
    try:
        uuid_obj = uuid.UUID(uuid_to_test, version=version)
    except ValueError:
        return False
    return str(uuid_obj) == uuid_to_test


def evalute_field(record, field_spec):
    """
    Evalute a field of a record using the type of the field_spec as a guide.
    """
    if type(field_spec) is int:
        return str(record[field_spec])
    elif type(field_spec) is str:
        return str(getattr(record, field_spec))
    else:
        return str(field_spec(record))


def table(records, fields, headings=None, alignment=None, file=None):
    """
    https[:]//stackoverflow[.]com/questions/13394140/generate-markdown-tables
    Generate a Doxygen-flavor Markdown table from records.

    file -- Any object with a 'write' method that takes a single string
        parameter.
    records -- Iterable.  Rows will be generated from this.
    fields -- List of fields for each row.  Each entry may be an integer,
        string or a function.  If the entry is an integer, it is assumed to be
        an index of each record.  If the entry is a string, it is assumed to be
        a field of each record.  If the entry is a function, it is called with
        the record and its return value is taken as the value of the field.
    headings -- List of column headings.
    alignment - List of pairs alignment characters.  The first of the pair
        specifies the alignment of the header, (Doxygen won't respect this, but
        it might look good, the second specifies the alignment of the cells in
        the column.

        Possible alignment characters are:
            '<' = Left align (default for cells)
            '>' = Right align
            '^' = Center (default for column headings)
    """

    num_columns = len(fields)
    if headings:
        assert len(headings) == num_columns

    # Compute the table cell data
    columns = [[] for i in range(num_columns)]
    for record in records:
        for i, field in enumerate(fields):
            columns[i].append(evalute_field(record, field))

    # Fill out any missing alignment characters.
    extended_align = alignment if alignment != None else []
    if len(extended_align) > num_columns:
        extended_align = extended_align[0:num_columns]
    elif len(extended_align) < num_columns:
        extended_align += [("^", "<") for i in range[num_columns - len(extended_align)]]

    heading_align, cell_align = (x for x in zip(*extended_align))  # noqa: B905

    field_widths = [
        len(max(column, key=len)) if len(column) > 0 else 0 for column in columns
    ]
    if headings:
        heading_widths = [max(len(head), 2) for head in headings]

    else:
        heading_widths = field_widths

    column_widths = [max(x) for x in zip(field_widths, heading_widths)]  # noqa: B905

    _ = " | ".join(
        [
            "{:" + a + str(w) + "}"
            for a, w in zip(heading_align, column_widths)  # noqa: B905
        ]
    )
    heading_template = "| " + _ + " |"
    _ = " | ".join(
        [
            "{:" + a + str(w) + "}"
            for a, w in zip(cell_align, column_widths)  # noqa: B905
        ]
    )

    row_template = "| " + _ + " |"

    _ = " | ".join(
        [
            left_rule[a] + "-" * (w - 2) + right_rule[a]
            for a, w in zip(cell_align, column_widths)  # noqa: B905
        ]
    )
    ruling = "| " + _ + " |"

    if file is not None:
        if headings:
            file.write(heading_template.format(*headings).rstrip() + "\n")
        file.write(ruling.rstrip() + "\n")
        for row in zip(*columns):  # noqa: B905
            file.write(row_template.format(*row).rstrip() + "\n")
        file.write("\n")
        file.write("\n")


def create_markdown_file(request, json_data, filename, fileuuid=""):
    """Создает markdown-файл из JSON-данных"""
    data_items = json_data.items()

    if request:
        update_callback_status(
            request, fileuuid, "git_upload_status", "Создание md файла"
        )

    try:
        file_path = os.path.join(
            settings.MEDIA_ROOT, f"{settings.MD_FILE_PATH}/{filename}.md"
        )
        with codecs.open(file_path, "w", encoding="utf-8") as file:
            func_map = {
                "acl_create_info.html": md_write_acl_create_info,
                "acl_internal_resources.html": md_write_acl_internal_resources,
                "acl_dmz_resources.html": md_write_acl_dmz_resources,
                "acl_external_resources.html": md_write_acl_external_resources,
                "acl_traffic.html": md_write_acl_traffic,
            }
            for key, data in data_items:
                if key in func_map:
                    func_map[key](file, data)

    except Exception as e:
        md_handle_error(filename, fileuuid, e)
        return False
    if request:
        update_callback_status(
            request, fileuuid, "git_upload_status", "The md file Created"
        )
    return file_path


def md_write_acl_create_info(file, data):
    logger.debug("[MakeMarkDown] acl_create_info")
    print("[MakeMarkDown] acl_create_info")
    file.write(f"## {data[4]}\n")
    file.write("##### Описание доступа к ресурсам\n")

    if data:
        fields = [0, 1]
        tmp = zip(contact_table, data)  # noqa: B905
        table(
            records=tmp,
            fields=fields,
            headings=contact_column,
            alignment=[("<", "<")] * len(contact_column),
            file=file,
        )


def md_write_acl_internal_resources(file, data):
    logger.debug(
        "[write_acl_internal_resources_to_markdown] Запись внутренних ресурсов"
    )
    print("[write_acl_internal_resources_to_markdown] Запись внутренних ресурсов")

    file.write("\n##### Список внутренних ресурсов (СГ АльфаСтрахование)\n")

    if data:
        table(
            records=data,
            fields=list(range(len(data[0])) if data else []),
            headings=standart_column,
            alignment=[("<", "<")] * len(standart_column),
            file=file,
        )


def md_write_acl_dmz_resources(file, data):
    logger.debug("[write_acl_dmz_resources_to_markdown] Запись DMZ ресурсов")
    print("[write_acl_dmz_resources_to_markdown] Запись DMZ ресурсов")

    file.write("\n##### Список DMZ ресурсов (СГ АльфаСтрахование)\n")

    if data:
        table(
            records=data,
            fields=list(range(len(data[0])) if data else []),
            headings=standart_column,
            alignment=[("<", "<")] * len(standart_column),
            file=file,
        )


def md_write_acl_external_resources(file, data):
    logger.debug("[write_acl_external_resources_to_markdown] Запись внешних ресурсов")
    print("[write_acl_external_resources_to_markdown] Запись внешних ресурсов")

    file.write("\n##### Список внешних ресурсов (Internet)\n")

    if data:
        table(
            records=data,
            fields=list(range(len(data[0])) if data else []),
            headings=external_column,
            alignment=[("<", "<")] * len(external_column),
            file=file,
        )


def md_write_acl_traffic(file, data):
    logger.debug("[write_acl_traffic_to_markdown] Запись потоков трафика")
    print("[write_acl_traffic_to_markdown] Запись потоков трафика")

    file.write("\n##### Потоки трафика\n")

    if data:
        # TODO временно до момента реструктуризации поля Acltext
        # max_length = max(len(item) for item in data)
        max_length = len(traffic_column)
        data = prepare_data_for_table(data, traffic_column)

        table(
            records=data,
            fields=list(range(max_length)),
            headings=traffic_column,
            alignment=[("<", "<")] * max_length,
            file=file,
        )


def prepare_data_for_table(data, fields):
    max_length = len(fields)
    prepared_data = []
    for item in data:
        prepared_item = item[:max_length]
        prepared_item.extend([""] * (max_length - len(prepared_item)))
        prepared_data.append(prepared_item)
    return prepared_data


def md_handle_error(filename, fileuuid, e):
    print(
        f"[ACL PORTAL] Error creating MD: {e} | filename: {filename} | fileuuid: {fileuuid}"
    )
    send_to_mattermost(
        f"[ACL PORTAL] Error creating MD: {e} | filename: {filename} | fileuuid: {fileuuid}"
    )
    logger.error(
        f"[ACL PORTAL] Error creating MD: {e} | filename: {filename} | fileuuid: {fileuuid}"
    )


class GitWorker:
    def __init__(
        self,
        request,
        GITPRO=None,
        USERNAME=None,
        PASSWORD=None,
        PATH_OF_GIT_REPO=None,
        MDFILE="",
        taskid="",
    ):
        os.environ["GIT_TIMEOUT"] = "60"
        uid = str(uuid.uuid4())
        self.USERNAME = settings.GITLAB_AUTH_USERNAME
        self.gitlab_token = settings.GIT_ACCESS_TOKEN
        self.USERNAME_encoded = urllib.parse.quote_plus(self.USERNAME)
        self.gitlab_token_encoded = urllib.parse.quote_plus(self.gitlab_token)
        if PATH_OF_GIT_REPO is not None:
            try:
                self.repo = git.Repo.init(
                    PATH_OF_GIT_REPO, bare=True
                )  # PATH_OF_GIT_REPO
            except:
                update_callback_status(
                    request,
                    taskid,
                    "git_upload_status",
                    f"Ошибка при инициализации Repo: {PATH_OF_GIT_REPO}",
                    0,
                )
                return False
            if settings.DEBUG:
                logger.debug(f"Инициализация GIT репозитория {PATH_OF_GIT_REPO}")
        else:
            tmp = os.path.join(tempfile.gettempdir(), uid)
            try:
                self.repo = git.Repo.init(
                    tmp, bare=True
                )  # uid, bare=True os.path.join(tempfile.gettempdir(), uid)
                print(f"Инициализирован новый репозиторий:{self.repo}")
            except:
                update_callback_status(
                    request,
                    taskid,
                    "git_upload_status",
                    f"Ошибка при инициализации Repo: {tmp}",
                    0,
                )
                return False
            if settings.DEBUG:
                logger.debug(
                    f"Инициализация GIT репозитория {os.path.join(tempfile.gettempdir(), uid)}"
                )

        self.request = request
        self.taskid = taskid
        update_callback_status(
            request, taskid, "git_upload_status", "Инициализация Git проекта"
        )

        # if PASSWORD is not None and USERNAME is not None:
        #      self.USERNAME = USERNAME
        #
        #      if '@' in self.USERNAME:
        #         self.USERNAME = self.USERNAME.replace('@', '%40')
        #      else:
        #          self.USERNAME = self.USERNAME + '%40' + 'alf'+'ast' + 'rah'+'.ru'
        #
        #      if PASSWORD:
        #          self.PASSWORD = urllib.parse.quote_plus(PASSWORD)
        #      # if '@' in self.PASSWORD:
        #      #     logger.warning('В пароле пользователя {} имеется запрещенный символ'.format(self.USERNAME))
        #

        self.GITURL = GITPRO
        self.GITPRO = GITPRO.split("://")[1]
        self.GITPRO = (
            f"https://{self.USERNAME_encoded}:{self.gitlab_token_encoded}@{self.GITPRO}"
        )

        if PATH_OF_GIT_REPO is not None:
            self.PATH_OF_GIT_REPO = PATH_OF_GIT_REPO
        else:
            # if settings.DEBUG:
            # self.PATH_OF_GIT_REPO = os.path.join(os.path.abspath(os.getcwd()), str(uuid.uuid4()))
            # else:
            self.PATH_OF_GIT_REPO = os.path.join(tempfile.gettempdir(), uid)

        self.PATH_OF_GIT_REPO = os.path.join(self.PATH_OF_GIT_REPO, "REPO")

        if not os.path.exists(self.PATH_OF_GIT_REPO):
            os.makedirs(self.PATH_OF_GIT_REPO)
            update_callback_status(
                request, taskid, "git_upload_status", "Создание временой папки"
            )
            # self.request.session['git_upload_status'].append({'status': "Создание временой папки: {}".format(self.PATH_OF_GIT_REPO)})
            if settings.DEBUG:
                logger.debug(f"Создание временой папки: {self.PATH_OF_GIT_REPO}")
        # else:
        # os.path.join(BASE_DIR, 'upload')
        if not os.path.exists(MDFILE):
            self.MDFILE = os.path.join(os.path.abspath(os.getcwd()), MDFILE)
        else:
            self.MDFILE = MDFILE
        if settings.DEBUG:
            logger.debug(f"Путь к md файлу: {self.MDFILE}")

    def free(self):
        for i in range(1, 3):
            self.repo.close()
            self.repo.__del__()
            if shutil.rmtree(Path(self.PATH_OF_GIT_REPO).parent, ignore_errors=True):
                break
            else:
                time.sleep(i)

    def clone(self):
        try:
            # if settings.DEBUG:
            #     logger.debug('Копируем репозиторий: {} ->{} '.format(self.GITPRO, self.PATH_OF_GIT_REPO))
            update_callback_status(
                self.request,
                self.taskid,
                "git_upload_status",
                "Клонируем удаленный репозиторий",
            )
            self.repo = self.repo.clone_from(self.GITPRO, self.PATH_OF_GIT_REPO)
        except Exception as e:
            if e.status == 128:
                # self.request.session['git_upload_status'].append({'error': "Нет доступа к GIT репозиторию"})
                update_callback_status(
                    self.request,
                    self.taskid,
                    "git_upload_status",
                    "Нет доступа к GIT репозиторию",
                    0,
                )
            else:
                # self.request.session['git_upload_status'].append({'error': "[Ошибка] {}".format(e)})
                update_callback_status(
                    self.request,
                    self.taskid,
                    "git_upload_status",
                    "Ошибка при клонировании репозитория",
                    0,
                )
            if settings.DEBUG:
                logger.debug("Ошибка при копировании")
            return 0

        if len(self.repo.index.entries) == 0:
            # self.request.session['git_upload_status'].append({'error': "Не удалось скачать файлы проекта, папка пустая"})
            update_callback_status(
                self.request,
                self.taskid,
                "git_upload_status",
                "Не удалось скачать файлы проекта, проект пустой",
                0,
            )
            if settings.DEBUG:
                logger.debug("Не удалось скачать файлы проекта, папка пустая")
            return 0
        # self.request.session['git_upload_status'].append({'status': "Скачано: {} файлов".format(len(self.repo.index.entries))})
        update_callback_status(
            self.request,
            self.taskid,
            "git_upload_status",
            f"Скачано: {len(self.repo.index.entries)} файлов",
        )
        if settings.DEBUG:
            logger.debug(f"Скачано: {len(self.repo.index.entries)} файлов")
        return True

    def fetch(self):
        try:
            self.repo.git.fetch()
            logger.info("Выполнен git fetch")
            return True
        except Exception as e:
            logger.error(f"Ошибка при выполнении git fetch:{e}")
            return False

    def pull(self):
        try:
            self.repo.git.pull()
            logger.info("git pull выполнен")
        except Exception as e:
            if e.status == 128:
                update_callback_status(
                    self.request,
                    self.taskid,
                    "git_upload_status",
                    "Нет доступа к GIT репозиторию",
                    0,
                )
            else:
                update_callback_status(
                    self.request,
                    self.taskid,
                    "git_upload_status",
                    "Ошибка при пулле резозитория",
                    0,
                )
                logger.error(e)

    def activity(self, git_filename: str):
        try:
            sfile = self.MDFILE
            dfile = os.path.join(self.PATH_OF_GIT_REPO, git_filename)
            if not copyfile(sfile, dfile):
                update_callback_status(
                    self.request,
                    self.taskid,
                    "git_upload_status",
                    "Ошибка при копировании md файла в проект",
                    0,
                )
                # self.request.session['git_upload_status'].append({'error': "Ошибка при копировании файла в проект: {}".format(dfile)})
                if settings.DEBUG:
                    logger.debug(
                        f"Ошибка при копировании файла {sfile} в проект: {dfile}"
                    )
                return 0
            update_callback_status(
                self.request,
                self.taskid,
                "git_upload_status",
                "Копирование md файла в проект",
            )
            # self.request.session['git_upload_status'].append({'status': "Копирование файла в проект: {}".format(dfile)})
            if settings.DEBUG:
                logger.debug(f"Копирование файла в проект: {dfile}")
        except Exception:
            # self.request.session['git_upload_status'].append({'error': "Возникла ошибка при копировании md файла в папку проекта"})
            update_callback_status(
                self.request,
                self.taskid,
                "git_upload_status",
                "Возникла ошибка при копировании md файла в папку проекта",
                0,
            )
            if settings.DEBUG:
                logger.error("Возникла ошибка при копировании md файла в папку проекта")
            return 0
        # finally:
        # if 'linux' in sys.platform:
        # return str(PurePosixPath(dfile)).replace('/', '//')
        return dfile  # str(PurePosixPath(dfile)).replace('/', '//')

    def addindex(self, filename):
        try:
            index = self.repo.index
            index.add([filename])
            clear_filename = os.path.basename(filename)
            username = self.USERNAME.replace("%40", "@")
            index.commit(
                f"{COMMIT_MESSAGE} {os.path.basename(clear_filename)} by {username}"
            )
            update_callback_status(
                self.request,
                self.taskid,
                "git_upload_status",
                "Локальный коммит изменений",
            )
            if settings.DEBUG:
                logger.debug("Локальный коммит изменений")
        except Exception as e:
            update_callback_status(
                self.request,
                self.taskid,
                "git_upload_status",
                "Ошибка при локальном коммите",
                0,
            )
            if settings.DEBUG:
                logger.error(f"Ошибка при локальном коммите: {e}")
            return False
        return True

    def push(self, refspec=""):
        if settings.DEBUG:
            logger.debug("Отправка изменений на сервер")
        try:
            for remote in self.repo.remotes:
                if remote.name == "origin":
                    if not remote.exists():
                        error_message = (
                            "Репозиторий не существует или нет доступа к нему"
                        )
                        print(error_message)
                        update_callback_status(
                            self.request,
                            self.taskid,
                            "git_upload_status",
                            error_message,
                            0,
                        )
                        return False
                    break
            for head in self.repo.heads:
                if head.name == refspec.split(":")[0]:
                    break
            else:
                error_message = f"Ветка {refspec.split(':')[0]} не существует"
                print(error_message)
                update_callback_status(
                    self.request, self.taskid, "git_upload_status", error_message, 0
                )
                return False

            print(f"repo remotes:{self.repo.remotes}")
            print(f"refspec:{refspec}")
            result = self.repo.remotes.origin.push(refspec=refspec, force=True)
            if result:
                print(f"Файл acl успешно загружен в репозиторий:{result}")
                update_callback_status(
                    self.request,
                    self.taskid,
                    "git_upload_status",
                    "Файл acl успешно загружен в репозиторий",
                )
                return True
        except git.exc.GitCommandError as e:
            if e.status == 128:
                error_message = "Ошибка аутентификации для данного репозитория"
                print(error_message)
                update_callback_status(
                    self.request, self.taskid, "git_upload_status", error_message, 0
                )
                send_to_mattermost(
                    f'[ACL PORTAL] Не удалось отправить md файл в git: У пользователя {self.USERNAME.replace("%40", "@")} нет доступа.'
                )
            else:
                error_message = "Ошибка при отправке коммита на сервер"
                print(error_message)
                update_callback_status(
                    self.request, self.taskid, "git_upload_status", error_message, 0
                )
                send_to_mattermost(
                    f"[ACL PORTAL] Не удалось отправить md файл в git: {e}"
                )
        finally:
            self.repo.close()
            if settings.DEBUG:
                print("Очистка временной папки")
                logger.debug("Очистка временной папки")
        return False


def dns_fileHandler(fname, full_buf) -> int:
    """Функция парсинга DNS Aktur"""
    count = 0
    EMAIL = "(\w|\.|\_|\-)+[@](\w|\_|\-|\.)+[.]\w{2,3}"
    DEFAULT = "Default"
    MASK = ["@", ";", " ", "_", "$"]
    with codecs.open(
        fname, "r", encoding="utf-8", errors="ignore"
    ) as f:  # encoding='utf-8'
        owner = ""

        for line in f:
            try:
                # line = line.decode('cp1252').encode('utf-8')
                if len(line) > 1:
                    if line[0] not in MASK:
                        if "IN" in str(line).upper():
                            if owner != "":
                                if owner not in full_buf:
                                    full_buf[owner] = []
                                full_buf[owner].append(line.strip())
                                count += 1
                            else:
                                if DEFAULT not in full_buf:
                                    full_buf[DEFAULT] = []
                                full_buf[DEFAULT].append(line.strip())
                                count += 1
                    else:
                        if (
                            ("<<" in line)
                            or (";<" in line and not ("/" in line or "\\" in line))
                            or ("<" in line and re.search(EMAIL, line))
                            or ("OT" in line and (":" in line or "#" in line))
                        ):
                            owner = line.strip()
                        elif (">>" in line and owner != "") or (
                            ";<" in line and ("/" in line or "\\" in line)
                        ):
                            if owner not in full_buf:
                                full_buf[owner] = []
                            owner = ""
                        else:
                            if DEFAULT not in full_buf:
                                full_buf[DEFAULT] = []
                            if line[0] not in MASK:
                                full_buf[DEFAULT].append(line.strip())
                                count += 1
            except:
                pass

    return count


def ExtractDataDns(uploaded_file_url) -> int:
    """Функция записи из буфера DNS файла"""
    buff = {}
    Tags = apps.get_model("ownerlist", "Tags")
    Iplist = apps.get_model("ownerlist", "Iplist")
    Owners = apps.get_model("ownerlist", "Owners")
    count = 0
    result = dns_fileHandler(uploaded_file_url, buff)
    file_name = os.path.basename(uploaded_file_url) or "Aktur DNS"
    if buff:
        for owner in buff:
            for value in buff[owner]:
                line = value.split()
                if len(line) < 4:
                    continue
                elif len(line) >= 5:
                    del line[1]
                if line[2] == "CNAME":
                    for cname in buff:
                        for tmp in buff[cname]:
                            s = tmp.split()
                            if line[3] == s[0]:
                                owner = f"{owner} CNAME {line[3]}"
                                if len(s) >= 5:
                                    line[3] = s[4]
                                else:
                                    line[3] = s[3]
                                break
                    if not isvalidip(line[3]):
                        try:
                            r = socket.gethostbyname(line[3])
                            if r:
                                owner = f"{owner} CNAME {line[3]}"
                                line[3] = r
                        except:
                            pass

                if settings.DEBUG:
                    print(line)

                try:
                    owner_info, created1 = Owners.objects.get_or_create(
                        username="Макаренко А.Б"
                    )

                except:
                    owner_info = Owners.get_default_owner()
                try:
                    tag_info, created2 = Tags.objects.get_or_create(name=file_name)
                except:
                    pass

                try:
                    created3 = None
                    ip_info, created3 = Iplist.objects.get_or_create(
                        ipv4=line[3],
                        hostname=line[0],
                        owner=owner_info,
                        comment=owner,
                    )
                    if created3:
                        count += 1
                except IntegrityError:
                    ip_info = Iplist.objects.get(ipv4=line[3])
                    ip_info.ipv4 = line[3]
                    ip_info.hostname = line[0]
                    ip_info.owner = owner_info
                    ip_info.comment = owner

                    ip_info.save()

                except DataError as e:
                    if settings.DEBUG:
                        print(f"- Ошибка данных: {e}")

                if created3:
                    ip_info.tags.add(tag_info)
                    ip_info.save()

    return count


def ClearSessionMeta(request=None):
    """Функция очистки сессии при переходе на другую страницу"""
    if request:
        SESSION_STORE = [
            "LOCAL_STORAGE",
            "uuid",
            "taskid",
            "ACT_MAKE_GIT",
            "ACT_MAKE_DOCX",
            "ACT_OMNI",
            "file_download",
        ]
        for sid in SESSION_STORE:
            if sid in request.session:
                del request.session[sid]

        if "file_download" in request.session:
            try:
                BASE = os.path.basename(request.session["file_download"])
                if BASE:
                    BASE = os.path.join(settings.BASE_DIR, "static//docx//" + BASE)
                    if os.path.exists(BASE):
                        os.remove(BASE)
            finally:
                del request.session["file_download"]
                BASE = None

        if "file_download_md" in request.session:
            try:
                BASE = os.path.basename(request.session["file_download_md"])
                if BASE:
                    BASE = os.path.join(settings.BASE_DIR, "static//md//" + BASE)
                    if os.path.exists(BASE):
                        os.remove(BASE)
            finally:
                del request.session["file_download_md"]
                BASE = None


def MakeTemporaryToken():
    s = f"{random.randrange(999)} ACL token {datetime.datetime.now()}"
    return f"{hashlib.md5(s.encode()).hexdigest()[:10]}"


def compare_first_column(doc1, doc2):
    table1 = doc1.tables[0]
    table2 = doc2.tables[0]
    column1 = [cell.text for cell in table1.columns[0].cells]
    column2 = [cell.text for cell in table2.columns[0].cells]
    if column1 == column2:
        return True
    return False


def compare_table_headers(doc1, doc2):
    for table1, table2 in zip(doc1.tables[1:], doc2.tables[1:]):  # noqa: B905
        headers1 = [cell.text for cell in table1.rows[0].cells]
        headers2 = [cell.text for cell in table2.rows[0].cells]
        if headers1 != headers2:
            return False
    return True


def check_file_against_template(file_path, template_path):
    doc1 = Document(file_path)
    doc2 = Document(template_path)
    return compare_first_column(doc1, doc2) and compare_table_headers(doc1, doc2)


def ParseDocx(file):
    template_file = os.path.join(BASE_DIR, "templates//ACL.docx")
    LOCAL_STORAGE = {}
    FIX_STATISTICS = 0
    SKIP__STATISTICS = 0
    result = {}
    if not check_file_against_template(file, template_file):
        logger.error(f"Был загружен файл, несоответствующий шаблону.{file}")
        result["error"] = (
            "Неправильный формат файла. Проверьте заголовки в вашем файле."
        )
        return result
    try:
        doc = Document(file)
        tables = doc.tables
        for idx, table in enumerate(tables):
            if idx >= len(FORM_APPLICATION_KEYS):
                continue
            table_rows = len(table.rows)
            table_cols = len(table.columns)
            namespace = FORM_APPLICATION_KEYS[idx]
            buffer = []
            if idx == 0:
                if table_rows > 7:
                    for id, row in enumerate(table.rows):
                        if table_rows >= 9 and id == 9:
                            continue
                        buffer.append(row.cells[1].text)
                else:
                    for id, row in enumerate(table.rows):
                        if id == 3 or id == 4:
                            buffer.append("")

                        elif id == 6 and row.cells[1].text == "":
                            buffer.append("Нет")
                            continue
                        elif id == 5 and row.cells[1].text == "" and buffer[6] != "":
                            buffer.append(buffer[6])
                            continue

                        if id in [4, 5, 6]:
                            if row.cells[1].text != "" and "." in row.cells[1].text:
                                if re.match("^(\d{2}.\d{2}.\d{4})$", row.cells[1].text):
                                    d = datetime.datetime.strptime(
                                        row.cells[1].text, "%d.%m.%Y"
                                    )
                                    buffer.append(d.strftime("%Y-%m-%d"))
                                    continue

                        buffer.append(row.cells[1].text)

                LOCAL_STORAGE[namespace] = buffer
            elif idx in range(len(FORM_APPLICATION_KEYS)):
                if table_cols >= 3:
                    # if namespace not in LOCAL_STORAGE:
                    # LOCAL_STORAGE[namespace] = []
                    for id, row in enumerate(table.rows):
                        if id != 0:
                            if any([cell.text != "" for cell in row.cells]):
                                line = [row.cells[x].text for x in range(table_cols)]
                                if len(line) == 3:
                                    CHECK_INDEX = 0
                                    if line[1] == "":
                                        line[1] = "32"
                                        FIX_STATISTICS += 1
                                    elif "." in line[1]:
                                        line[1] = GetNumberFromMask(line[1])
                                    elif "/" in line[1]:
                                        line[1] = line[1][1:]
                                elif len(line) == 4:
                                    if line[2] == "":
                                        line[2] = "32"
                                        FIX_STATISTICS += 1

                                if namespace not in LOCAL_STORAGE:
                                    LOCAL_STORAGE[namespace] = []
                                LOCAL_STORAGE[namespace].append(line[:])
                            else:
                                SKIP__STATISTICS += 1

    except ValueError as e:
        if "is not a Word" in str(e):
            result["error"] = "Ошибка при загрузке старого формата doc файла"
        else:
            result["error"] = f"{e}"
    except opc.exceptions.PackageNotFoundError:
        result["error"] = "Ошибка, вероятно это не docx файл"
    except Exception as e:
        result["error"] = str(e)
        logger.error(f"Ошибка загрузки из файла:{e}")
    finally:
        if len(LOCAL_STORAGE) > 0:
            result["LOCAL_STORAGE"] = LOCAL_STORAGE
        result["META"] = {}
        if FIX_STATISTICS > 0:
            result["META"]["fixed"] = f"{FIX_STATISTICS}"
        if SKIP__STATISTICS > 0:
            result["META"]["skipped"] = f"{SKIP__STATISTICS}"
    return result


def GetNumberFromMask(mask=""):
    """Функция возвращает число префикса по маске"""
    result = mask
    buff = [
        "000.000.000.000",
        "128.000.000.000",
        "192.000.000.000",
        "224.000.000.000",
        "240.000.000.000",
        "248.000.000.000",
        "252.000.000.000",
        "254.000.000.000",
        "255.000.000.000",
        "255.128.000.000",
        "255.192.000.000",
        "255.224.000.000",
        "255.240.000.000",
        "255.248.000.000",
        "255.252.000.000",
        "255.254.000.000",
        "255.255.000.000",
        "255.255.128.000",
        "255.255.192.000",
        "255.255.224.000",
        "255.255.240.000",
        "255.255.248.000",
        "255.255.252.000",
        "255.255.254.000",
        "255.255.255.000",
        "255.255.255.128",
        "255.255.255.192",
        "255.255.255.224",
        "255.255.255.240",
        "255.255.255.248",
        "255.255.255.252",
        "255.255.255.254",
        "255.255.255.255",
    ]
    try:
        result = buff.index(str(mask))
    except:
        return mask
    return str(result)


def send_onmitracker(
    sender, title, text, attach, proxy=None, fake=None, request=None, uid=None
):
    """Функция для отправки обращения напрямую в SD получения номера зарегистрированного обращения"""

    try:
        send_to_mattermost(
            f"Попытка отправки обращения в omnitracker. sender: {sender}; title:{title}; Ссылка на ACL :https://acl.vesta.ru/acl/info/{str(uid)}"
        )
    except Exception as e:
        logger.warning(f"[SEND_TO_MATTERMOST EXCEPTION] {e}")
    bodyNotEncoded = f"""<?xml version="1.0" encoding="utf-8"?>
                <soap:Envelope xmlns:soap="http://schemas.xmlsoap.org/soap/envelope/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance" xmlns:s="http://www.w3.org/2001/XMLSchema" xmlns:tns="http://www.omninet.de/OtWebSvc/v1">
        <soap:Body>
            <tns:InvokeScript>
               <tns:Script name="OmniAPI" runAt="Server">
                    <tns:Parameters>
                        <tns:StringVal name='strFunction'>CreateIncidentFromTemplateWithAttachmentAndExternalNumber</tns:StringVal>
                        <tns:StringVal name='strEMail'>{sender}</tns:StringVal>
                        <tns:StringVal name='strTitle'>{title}</tns:StringVal>
                        <tns:StringVal name='strDescription'>{text}</tns:StringVal>
                        <tns:StringVal name='strTemplateCode'>SH0458</tns:StringVal>
                        <tns:StringVal name='strExternalNumber'>{str(uid)}</tns:StringVal>
                        <tns:StringVal name='strAttachmentPath'>{attach}</tns:StringVal>
                    </tns:Parameters>
                </tns:Script>
            </tns:InvokeScript>
        </soap:Body>
        </soap:Envelope>
        """

    body = bodyNotEncoded.encode("utf-8")
    logger.debug("BODY NOT ENCODED:" + str(bodyNotEncoded))
    logger.debug("BODY ENCODED:" + str(body))
    logger.debug("START SEND OMNI")

    try:
        auth = HttpNtlmAuth(
            f"{settings.LDAP_AUTH_CONNECTION_DOMAIN}\\{settings.LDAP_AUTH_CONNECTION_USERNAME}",
            settings.LDAP_AUTH_CONNECTION_PASSWORD,
        )
    except Exception as e:
        send_to_mattermost(f"Ошибка авторизации LDAP: {e}")
    result = 0

    for i in range(settings.ATTEMPS):
        try:
            # time.sleep(5)
            logger.debug(f"[{i}] Отправка запроса в {settings.OMNITRACKER_URL}")
            logger.debug("[TRACE] Тело запроса: " + body.decode())
            if settings.DEBUG:
                res = mock_send_omnitracker_200(
                    taskid="777777",
                    omnitracker_url=settings.OMNITRACKER_URL,
                    body=body,
                    headers=OMNI_HTTP_HEADERS,
                    timeout=settings.TIMEOUT,
                    auth=auth,
                )
            else:
                res = requests.post(
                    settings.OMNITRACKER_URL,
                    body,
                    headers=OMNI_HTTP_HEADERS,
                    timeout=settings.TIMEOUT,
                    auth=auth,
                )
            if i == 2:
                send_to_mattermost(
                    f"Попытка подключения номер: {i}.",
                    f"res: {res.text}",
                    f"Ссылка на ACL: https://acl.vesta.ru/acl/info/{str(uid)}",
                )
            if res and res.status_code == 200:
                result = res.text
                logger.info("OMNITRACKER ANSWER ====" + result)
                if result and len(result) > 3 and "strResponse" in result:
                    result = result.split('strResponse">')
                    if len(result) > 1:
                        result = result[1].split("</")[0]
                    else:
                        result = res.text.split("strResponse>")
                        if len(result) > 1:
                            result = result[1].split("</")[0]
                        else:
                            update_callback_status(
                                request,
                                uid,
                                "omni_email_status",
                                "Сервер omni вернул неправильный ответ",
                                0,
                            )
                            send_to_mattermost(
                                f"[TRACE] Сервер omni вернул неправильный ответ: {res.text}"
                                f"Попытка номер: {i}",
                                f"Ссылка на ACL :https://acl.vesta.ru/acl/info/{str(uid)}",
                            )
                            if settings.DEBUG:
                                logger.error(
                                    f"[TRACE] Сервер omni вернул неправильный ответ {res.text} "
                                )

                            result = 0
                break
            else:
                if res.status_code == 401 and not auth:
                    if "requests_ntlm" in sys.modules:
                        try:
                            auth = HttpNtlmAuth(
                                f"{settings.LDAP_AUTH_CONNECTION_DOMAIN}\\{settings.LDAP_AUTH_CONNECTION_USERNAME}",
                                settings.LDAP_AUTH_CONNECTION_PASSWORD,
                            )
                        except Exception as e:
                            send_to_mattermost(f"Ошибка авторизации LDAP: {e}")
                        logger.info("[TRACE] Пытаемся авторизоватся в omni")
                    else:
                        logger.error("Нужный модуль requests_ntlm не загружен :-(")
                if i >= settings.ATTEMPS:
                    break
                update_callback_status(
                    request,
                    uid,
                    "omni_email_status",
                    f"Повторная отправка сообщения в SD ({i})",
                    0,
                )
                sleep(i)
        except requests.Timeout:
            # if settings.DEBUG:
            send_to_mattermost(
                "Подключение не удалось по таймауту (Timeout).",
                f"Попытка подключения номер: {i}.",
                f"OMNITRACKER_URL: {settings.OMNITRACKER_URL}",
                f"Ссылка на ACL: https://acl.vesta.ru/acl/info/{str(uid)}",
            )
            logger.error(
                f"[TRACE] Подключение не удалось по таймауту к {settings.OMNITRACKER_URL}"
            )
            update_callback_status(
                request,
                uid,
                "omni_email_status",
                f"Повторная отправка сообщения в SD ({i})",
                1,
            )
        except requests.ConnectionError:
            # if settings.DEBUG:
            send_to_mattermost(
                "Сервис не отвечает (Connection Error).",
                f"Попытка подключения номер: {i}.",
                f"OMNITRACKER_URL: {settings.OMNITRACKER_URL}",
                f"Ссылка на ACL: https://acl.vesta.ru/acl/info/{str(uid)}",
            )
            logger.error(f"[TRACE] Сервис {settings.OMNITRACKER_URL} не отвечает")
            update_callback_status(
                request,
                uid,
                "omni_email_status",
                f"Повторная отправка сообщения в SD ({i})",
                1,
            )

        except Exception as e:
            send_to_mattermost(
                f"Ошибка Omni request: {e}.",
                f"Попытка номер: {i}.",
                f"OMNITRACKER_URL: {settings.OMNITRACKER_URL}",
                f"Ссылка на ACL: https://acl.vesta.ru/acl/info/ {str(uid)}",
            )
            update_callback_status(
                request,
                uid,
                "omni_email_status",
                f"Повторная отправка сообщения в SD ({i})",
                1,
            )
            logger.error("[Omni request] " + str(e))
            if settings.DEBUG:
                print(e)
    if not result:
        return 0
    send_to_mattermost(
        f"Обращение [https://acl.vesta.ru/acl/info/{str(uid)}] успешно отправлено. Получен ответ: [result:{result}]. Отправитель:{sender}"
    )
    return result


def omni_check_status(id=None, proxy=None):
    body = f"""<?xml version="1.0" encoding="utf-8"?>
<soap:Envelope xmlns:soap="http://schemas.xmlsoap.org/soap/envelope/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance" xmlns:s="http://www.w3.org/2001/XMLSchema" xmlns:tns="http://www.omninet.de/OtWebSvc/v1">
<soap:Body>
   <tns:InvokeScript>
     <tns:Script name="OmniAPI" runAt="Server">
                               <tns:Parameters>
                                               <tns:StringVal name="strFunction">GetIncStateByNumber</tns:StringVal>
                                               <tns:LongIntVal name="lngNumber">{id}</tns:LongIntVal>
                              </tns:Parameters>
    </tns:Script>
   </tns:InvokeScript>
</soap:Body>
</soap:Envelope>
""".encode()

    result = ""
    max_retries = 3
    try:
        auth = HttpNtlmAuth(
            f"{settings.LDAP_AUTH_CONNECTION_DOMAIN}\\{settings.LDAP_AUTH_CONNECTION_USERNAME}",
            settings.LDAP_AUTH_CONNECTION_PASSWORD,
        )
    except Exception as e:
        logger.error(f"Ошибка авторизации LDAP: {e}")
        print(f"Ошибка авторизации LDAP: {e}")

    for attempt in range(max_retries):
        try:
            if settings.DEBUG:
                r = mock_omni_check_status_200(
                    acl_status="Assigned",
                    omnitracker_url=settings.OMNITRACKER_URL,
                    body=body,
                    headers=OMNI_HTTP_HEADERS,
                    proxies=proxy,
                    auth=auth,
                    timeout=1,
                )
            else:
                r = requests.post(
                    settings.OMNITRACKER_URL,
                    body,
                    headers=OMNI_HTTP_HEADERS,
                    proxies=proxy,
                    auth=auth,
                    timeout=1,
                )
            if r.status_code == 200:
                try:
                    result = r.text
                    if result and len(result) > 3:
                        result = result.split('strResponse">')
                        result = result[1].split("</")[0]
                    return OMNI_ACL_STATUS.get(result), OMNI_HTTP_STATUS.get(result)
                except Exception as e:
                    logger.error("[Check status] " + str(e))
                    if settings.DEBUG:
                        print(e)
                    return None, None

        except Exception as e:
            logger.info(f"Ошибка на попытке {attempt + 1} ID:{id} - {e}")
            send_to_mattermost(
                f"[omni_check_status] Ошибка на попытке {attempt + 1} ID:{id} - {e}"
            )
            print(f"[omni_check_status] Ошибка на попытке {attempt + 1} ID:{id} - {e}")

            if attempt < max_retries - 1:
                time.sleep(15)
            else:
                return None, None


def SendMessageToApprove(acl_id, acl_owner, user, token=""):
    EMAIL_APPROVE = """<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.0 Transitional//EN" "http://www.w3.org/TR/xhtml1/DTD/xhtml1-transitional.dtd">
    <html xmlns="http://www.w3.org/1999/xhtml"><head>
    <meta charset="utf-8"><meta http-equiv="Content-Type" content="text/html; charset=UTF-8" />
    <meta name="viewport" content="width=device-width, initial-scale=1.0"/>
    <title>Согласование обращения</title>
    <style>.container{margin: 0 auto;padding: 10px; width: 600px;height: 200px;border-left: 3px solid #6a68d9;border-right: 1px solid #ccc;border-top: 1px solid #ccc;border-bottom: 1px solid #ccc;}.header{border-bottom: 1px solid #ccc;padding: 5px;margin: 15px;display: flex;}.header_text{color: #6a68d9;padding-right: 40px;display: inline;}.header_text_portal{ color: #484848; font-weight: bold;}
    .logo{display: block; float: left;}.container_body{margin: 15px;min-height: 165px;}
    .container_footer{display: flex;align-items: center;}
    .footer{color: #48484859;padding: 15px;border-top: 1px solid #ccc;max-height: 57px;}</style></head>
    <body>
    <div class="container">
    <table cellpadding="0" cellspacing="0" border="0">
    <tr style="border-bottom: 1px solid #ccc;">
    <td style="padding-right: 15px"><img src="data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAACIAAAAdCAYAAADPa766AAAAAXNSR0IArs4c6QAAAARnQU1BAACxjwv8YQUAAAAJcEhZcwAADsMAAA7DAcdvqGQAAAjmSURBVFhHnZf5cxTXEcf5P/NLChNwQAJZIAQ6AxEORcrEFQgkwcE2ijGC4IAx6OQQh2PMoZtokdBKitCFrtXu7Nz7zaffSjIkdlUqszQz86bf6+/r/na/1jYlkkkaS3FsLysqKa+UX5LwXtqUFKUSgsr/IzbfSUlxyFAJO6WCSnFJJd63/aCMkVLMzUCE3EtKGU9BmKa8pxGS/IfYN5OfekcwXn62+aESnlkaG6vc15VEEY5IAeJQ4A6hCDSHB/F9qZC3L+yAb/8t4f8gEc6Ot6Q8L1UR22kpYGUPW0X2XzAgjIrBkqc4AYgBRPwiX3BZCMaQ94D7ptiYeyZi77xviJvzE+8RUgycExTxUir5RMwBsT0buiIowV0oaXQ41Nk/dKu66jPt3XMT6UI6kXbtrbilShurKEvlnlvc21W5275tio3Zt29UsZu7fWONqsoeVe1r1/Wvs8qtpoqdZyLABAYEsuC4FM+kkOLJ/aKaajv1+acP1d05op6uLDKpbpPurJPOjgl1dmZ1796MOtondKdnWh23Xjm9nu6ybo/TnSjP4b2rYxJd5NaEjh+/r1MfTWtpHs9AgySJtY07QGyACMaxzpzI6XLrkvLww9wYohlCnGBDQgjEHtzdCWPFgDFcbaEKzeWmY99MD9l6N2G9V9lEVb8c0IO7oSLmRDFkNRBG14SwRKRv1fZhjfQVt2KaMBmeI0Y8WC8ffRawDLBvBg7rlmEWd9th6vSNoIHTT0gC26izgX4O/h2tn9BXbbMie908gCTuF6McwdLqHS81OuiDPlUhzLtakgIwgchJyqIJ5OIZW+WyQJxjP9HYiyWFHoB4D4ukPhaSxJLAdBl0tYh/SAEwdbUDunQxK59aEvMNIJZQpBbsDWBP9c4JDfX5LtGCdI2JBtniZwsZuwAVmSEeeQWXvJVQ9ftb9aB7TrHHOLbNzSU2UHJuQsxlrAUGFdhccyNALo3zvIxtOOLiYkUHF4aQpvK9ab0YDlwuuepK/KwgRVgIYkNPKJjjAeT1m1RTMwU9f17QgaoJ7fr5oHrvLMMZ88S6SgFsAiiRxARrpoQcJD5gmuv6dfnSFNzZCI3bLR8t7gFIK9/71xYQF2m8FDMes6MIQEVQzM5JF/86p6ampzp46IH2f5DR9p+9Vs3egpoanqBnxM+xLusEEMIIGXmuiPnUdx/DzXWDugIQ44g5awOI/YNaQKvckdWLEd+CBRCrtjE7zAEEEKz56P4bVb3frd+2ZHT96ooms9L0y1Q1e+Z1aN+4MhkfsLif8mnHRETsghBe8RzHEB0wgQPSD5DJt4DgYgNjTA+gfOUvxgBSZLicfAG+K8QzZFFR7dfmVV81oEe3lylIzGGRfCGQtxTrQxbODEXKA6DImlY9B4aKml8k5HAlZe0U75Yg0BaQtmwZCFIGYsTjxRQqd44ABC7wM96Yl0LOg/6+vGr3PNe9b1YVEfcYtydpDrDo4uvVRauUuJ1T1c6S3t68du/p0anfv9TKCus7ruNlNmYcKwOZAIgdjKWN0JhHAOLh0n3vD2jEAYEhVuRQKsKuj44/Vuu5ZXmE3roDS+mUhVMDxMoRKZw4Qib6/qmnquoeXWgtqA4u3Li2KI+jw1ItBYgH0CONI2r7clxeUm45toAkBKpoZN31mN2vu4JmbPa5Z8elw9X3NDnOiQrikAochQDlbkD8ArDRi+DQwDNfhw8+UtuVaS2vS385P6qWpn7lbQNWh8wOuo21Q2TNK4CsOe+W09d5BI6Qxru239XI8LojptUKi/W1v3k60TKi/DogiLVlRZgUmJMjpddc0YvRywwkaql7pc8/WYArFERAD42uqv7AsGZnCbFlE3bWqDVH68cAkuUdj5DjACmzxcqvRx050jCi3tt5+SiHgLF4nmh5qc/+OCcfY+tFX88GXyjv5+EUpGZ6gazIZIqqqxnU+T9NyTKWEsL+Ii2t5VS966my2VDrKK+xxsJyqsP7M+psh28AiWlBNoC40sXJEOvKl2uqO/idnvwj71xNUujD5nGdOzNVLmILsT6ouKu2L+bdWbSyJq3lqJTN3+vMWcLBc0TJjdlBQB7H1I36vRmABsqhP4/+hS9eqqFmWmMZOwyxA0+2Wdl1ldWICdHWlku6emlMRxvb1HT4M+fSc6eW9efTWUcyi2/fU6nhAGRrXdE/h1K1NGd0+uSAchgJMZzSeQVwCApp5rV0oKJf09lYt7rmVV1zUb/7+IaePcabJEdCWKyYloHwWAJMbGeI8YX72lIkbz2gIpbU8bWvloYBrS5T3FDwmfP4u4IO7hvFSAbQo1qcJobUiZR2zicuVgqslP/9xpSOND/R8gLZkkupP6FyOWsR4VtIrXL2LX35z3Vpls/GFXZh55sdaL4fUUNWtLiQqvngcz3soU4Q54ATM8d50t0zo7NnRvSGOmHtZQJBXVm36omBhdVYvzlxR+c+eeaqbZgU8cCSO8kNwKbEuO4HIK7dKXfYBsayyZ7ZAxNT3by6qMbqjKZIZfogGssVK3f8IHDqMdM2Y1wrUIk9F8Leh8vaf+C6JqbtkLR+htNci86eXdYRBmzI7m95xFpFQ2p/zwDJvAKYiAV8wrM4l+pYQ7/OnMxqDt7Ywh4APEAWiXMRl1ij5NGq5QrSg0cF1dZ2q6NrhkPOjg/aCrIrsfAxx4yXw2Ibfis0mISwuNUaH7BvKvqhtU3lehJa/jd8q/0VXep7EssjqzzK/TrjFFaM8UyLeeVyXjt3dNBvTLuQ5Kj51snZQtQ011DZ5eqPMZprA4iZsoG3xcYAZ6Ghy7KDy3Y8Oxvpq8uT+hUd1q/rRvXpuTFdvPBCra19On16QIdrnpPur/T4W18e4COyEYpjhTUNgC1bdsQ71wYQ0/gxsZDRrUFQq7wxh591U5Z2q29oCXpDfXzythoP3dSxY906f35Qw1TXiHQJmZMr5liFg9G1iyznQmGbtLXfvQCyeTnNd4WJlg2uS8S19geYHWpWyOweUhFDeGWduFVdmjlAEyLSMko8ZIlV7KxgLfOE3X/UJdK/AfklVKxVmjCRAAAAAElFTkSuQmCC" class="logo"></td>
    <td style="padding-right: 15px"><h4 class="header_text">Cогласование</h4></td>
    <td><div class="header_text_portal" style="padding-top: 10px"><h3 style="margin-left: 15px; color: #484848cc;">ACL Портал</h3></div></td> </tr> 
    <tr style="padding-top: 10px"> <td colspan="3" style="">Ожидается Ваше согласование</td> </tr> 
    <tr style="padding-top: 10px"> <td colspan="2" style="padding-top: 10px">Запрос от:</td> 
    <td style="font-weight: bold;">%s</td> </tr><tr style="padding-top: 10px"><td colspan="2">Подробнее:</td> 
    <td><a href="%s" style="color:#1a73e8;">Перейти на портал</a></td> </tr></table><div class="footer">
    <p>Данное сообщение сформировано автоматически порталом acl.vesta.ru, просьба не отвечать на него.</p><div> 
    </div>
    </body>
    </html>
                                    """ % (
        acl_owner,
        f"https://acl.vesta.ru/acl/pending/{acl_id}/?token={token}",
    )

    e = EmailMessage(
        subject="Согласование обращения " + str(user.username),
        body=EMAIL_APPROVE,
        from_email="acl@alfastrah.ru",
        to=[user.email],
    )
    e.content_subtype = "html"
    e.send(fail_silently=settings.DEBUG)
    if settings.DEBUG:
        print(f"https://acl.vesta.ru/acl/pending/{acl_id}/?token={token}")


def send_to_mattermost(*messages):
    try:
        text = ""
        for message in messages:
            if settings.DEBUG:
                text += "[ACL DEBUG] "
            text += f"{message}\n"

        payload = {"text": text}
        response = requests.post(settings.MATTERMOST_WEBHOOK_URL, json=payload)
        logger.info(f"response = {response}")
        if response.status_code != 200:
            logger.warning(
                "[Send to Mattermost] Failed to send message. Error: ",
                response.status_code,
            )
    except Exception as e:
        logger.error(e)


def get_gitlab_project_info(session, repo_url: str, mode: str = ""):
    try:
        logger.info("[GET_GITLAB_PROJECT_INFO] Получение информации о проекте.")
        # base_api_url = f"https://{request.session['GIT_USERNAME']}%40alfastrah.ru:{request.session['GIT_PASSWORD']}@gitlab.alfastrah.ru/api/v4/projects"
        base_api_url = "https://gitlab.alfastrah.ru/api/v4/projects"
        project_namespace = "/".join(repo_url.split("/")[3:-1]).replace("/", "%2F")
        logger.info(
            f"[GET_GITLAB_PROJECT_INFO] Получен project namespace:{project_namespace}"
        )
        repo_url_parts = repo_url.split("/")
        project_name_with_extension = repo_url_parts[-1]
        project_name = os.path.splitext(project_name_with_extension)[0]
        logger.info(f"[GET_GITLAB_PROJECT_INFO] Получен project_name:{project_name}")
        project_id = None
        get_project_id_api = f"{base_api_url}/{project_namespace}%2F{project_name}"
        response = session.get(get_project_id_api)
        if response.status_code == 200:
            project_id = int(response.json()["id"]) or None

        if project_id is not None:
            logger.info(f"[GET_GITLAB_PROJECT_INFO] Получен project_id:{project_id}")
            if mode == "id":
                return project_id
            info = {
                "namespace": project_namespace,
                "project_name": project_name,
                "project_id": project_id,
            }
            return info
    except Exception as e:
        logger.error(f"Ошибка при получении информации gitlab: {e}")


def get_files_from_gitlab(repo_url: str, branch_name: str = "develop"):
    file_list = []
    try:
        session = requests.Session()
        session.headers.update({"PRIVATE-TOKEN": settings.GIT_ACCESS_TOKEN})
        project_id = get_gitlab_project_info(session, repo_url)["project_id"]
        base_api_url = "https://gitlab.alfastrah.ru/api/v4/projects"
        get_file_list_api = (
            f"{base_api_url}/{project_id}/repository/tree?ref={branch_name}"
        )
        response = session.get(get_file_list_api)
        if response.status_code == 200:
            file_list = [
                file_name["name"]
                for file_name in response.json()
                if ".md" in file_name["name"]
            ]
            logger.info("[GET_FILES_FROM_GITLAB] Список файлов успешно получен")
    except Exception as e:
        logger.error(f"Ошибка при получении списка файлов:{e}")
    finally:
        return file_list


def get_acl_from_gitlab(
    repo_url: str, branch_name: str = "develop", file_name: str = ""
):
    """Получает acl в формате str из gitlab"""
    try:
        logger.info("[GET_ACL_FROM_GITLAB] Получаю acl в формате str из gitlab")
        session = requests.Session()
        session.headers.update({"PRIVATE-TOKEN": settings.GIT_ACCESS_TOKEN})
        project_id = get_gitlab_project_info(session, repo_url)["project_id"]
        get_file_content_api = f"https://gitlab.alfastrah.ru/api/v4/projects/{project_id}/repository/files/{file_name}?ref={branch_name}"
        response = session.get(get_file_content_api)
        if response.status_code == 200:
            logger.info(
                "[GET_ACL_FROM_GITLAB] acl получен из gitlab. Обработка к необходимому формату."
            )
            encoded_file_content = response.json()["content"]
            decoded_file_content = base64.b64decode(encoded_file_content).decode(
                "utf-8"
            )
            project_desc = get_project_desc_from_git(
                session=session, project_id=str(project_id)
            )
            return decoded_file_content, project_desc
        logger.info("[GET_ACL_FROM_GITLAB] Не удалось получить acl из gitlab")
        return
    except Exception as e:
        logger.error(f"[ПОЛУЧЕНИЕ md_content] Ошибка: {e}")


def convert_md_to_dict(md_content: str):
    try:
        LOCAL_STORAGE = {}
        result = {}
        lines = md_content.split("\n")

        current_section = None
        for _, line in enumerate(lines):
            line = line.strip()
            if line.startswith("##### Описание доступа к ресурсам"):
                current_section = "acl_create_info.html"
                LOCAL_STORAGE[current_section] = []

            elif line.startswith("##### Список внутренних ресурсов"):
                current_section = "acl_internal_resources.html"
                LOCAL_STORAGE[current_section] = []

            elif line.startswith("##### Список DMZ ресурсов"):
                current_section = "acl_dmz_resources.html"
                LOCAL_STORAGE[current_section] = []

            elif line.startswith("##### Список внешних ресурсов"):
                current_section = "acl_external_resources.html"
                LOCAL_STORAGE[current_section] = []

            elif line.startswith("##### Потоки трафика"):
                current_section = "acl_traffic.html"
                LOCAL_STORAGE[current_section] = []

            elif line.startswith("|"):
                if current_section == "acl_create_info.html":
                    parts = [p.strip() for p in line.split("|")[1:-1]]
                    if re.match("^:--", parts[0]) or parts == contact_column:
                        continue
                    LOCAL_STORAGE[current_section].append(parts[1])

                elif current_section in [
                    "acl_internal_resources.html",
                    "acl_dmz_resources.html",
                    "acl_external_resources.html",
                    "acl_traffic.html",
                ]:
                    parts = [p.strip() for p in line.split("|")[1:-1]]
                    for part_idx, part in enumerate(parts):
                        if part == "":
                            parts[part_idx] = " "

                    if (
                        len(parts) > 0
                        and re.match("^:--", parts[0])
                        or parts == standart_column
                        or parts == external_column
                        or set(parts).issubset(set(traffic_column))
                    ):
                        continue
                    LOCAL_STORAGE[current_section].append(parts)

        result["LOCAL_STORAGE"] = LOCAL_STORAGE
        return result
    except Exception as e:
        logger.error(e)


def acl_sending_retry_checking(acl_object):
    """Проверяет acl в omnitracker на наличие заявки"""
    acl_taskid = acl_object.taskid
    acl_status = acl_object.status

    omni_taskid = check_taskId_by_uuid(acl_object.id)

    if omni_taskid:
        omni_acl_status, omni_http_status = omni_check_status(omni_taskid)
        if acl_taskid != omni_taskid:
            acl_object.taskid = str(omni_taskid)
            acl_object.save(update_fields=["taskid"])
        if acl_status != omni_acl_status and omni_acl_status in ["CMP", "JOB", "CNL"]:
            acl_object.status = omni_acl_status
            acl_object.save(update_fields=["status"])
        return omni_taskid
    else:
        return None


def check_taskId_by_uuid(uuid, proxy=None):
    """Функция проверяет номер SD по номеру UUID в omnitracker"""

    body = f"""<?xml version="1.0" encoding="utf-8"?>
    <soap:Envelope xmlns:soap="http://schemas.xmlsoap.org/soap/envelope/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance" xmlns:s="http://www.w3.org/2001/XMLSchema" xmlns:tns="http://www.omninet.de/OtWebSvc/v1">
      <soap:Body>
        <tns:InvokeScript>
          <tns:Script name="OmniAPI" runAt="Server">
            <tns:Parameters>
              <tns:StringVal name="strFunction">GetIncInfoByExternalNumber</tns:StringVal>
              <tns:StringVal name="strExternalNumber">{str(uuid)}</tns:StringVal>
            </tns:Parameters>
          </tns:Script>
        </tns:InvokeScript>
      </soap:Body>
    </soap:Envelope>
    """.encode()
    try:
        auth = HttpNtlmAuth(
            f"{settings.LDAP_AUTH_CONNECTION_DOMAIN}\\{settings.LDAP_AUTH_CONNECTION_USERNAME}",
            settings.LDAP_AUTH_CONNECTION_PASSWORD,
        )
    except Exception as e:
        logger.error(f"Ошибка авторизации LDAP: {e}")
    try:
        response = requests.post(
            settings.OMNITRACKER_URL,
            body,
            headers=OMNI_HTTP_HEADERS,
            proxies=proxy,
            auth=auth,
        )
        if response and response.status_code == 200:

            result = response.text
            logger.info("OMNITRACKER ANSWER ====" + result)

            if result and len(result) > 3:
                result = result.split('strResponse">')
                if len(result) > 1:
                    result = result[1].split("</")[0]
                else:
                    result = response.text.split("strResponse>")
                    if len(result) > 1:
                        result = result[1].split("</")[0]
                    else:
                        logger.warning(
                            f"uuid: {uuid}. Сервер omni вернул неправильный ответ"
                        )
                        result = ""
    except Exception as e:
        logger.error(f"[check_taskId_by_uuid] Error: {e}")
        result = ""

    try:
        result = int(result)
    except:
        result = ""
    return str(result)


def sync_acl_portal_projects_list(group_name: str = "ACL"):
    project_list = []
    session = requests.Session()
    session.headers.update({"PRIVATE-TOKEN": settings.GIT_ACCESS_TOKEN})

    def get_subgroup_projects(group_id, path=""):
        try:
            url_get_projects = (
                f"https://gitlab.alfastrah.ru/api/v4/groups/{group_id}/projects"
            )
            response = session.get(url_get_projects)
            if response.status_code == 200:
                for project in response.json():
                    full_path = f"{path}/{project['name']}".lstrip("/")
                    project_list.append(
                        {
                            "full_path": full_path,
                            "http_url_to_repo": project["http_url_to_repo"],
                        }
                    )
        except Exception as e:
            logger.error(f"Ошибка при получении проектов из подгруппы: {e}")

        try:
            url_get_subgroups = (
                f"https://gitlab.alfastrah.ru/api/v4/groups/{group_id}/subgroups"
            )
            response = session.get(url_get_subgroups)
            if response.status_code == 200:
                for subgroup in response.json():
                    subgroup_path = f"{path}/{subgroup['name']}".lstrip("/")
                    get_subgroup_projects(subgroup["id"], path=subgroup_path)
        except Exception as e:
            logger.error(f"Ошибка при получении подгрупп из подгруппы: {e}")

    try:
        url_get_subgroups = (
            f"https://gitlab.alfastrah.ru/api/v4/groups/{group_name}/subgroups"
        )
        response = session.get(url_get_subgroups)
        if response.status_code == 200:
            for block in response.json():
                block_path = f"{group_name}/{block['name']}".lstrip("/")
                get_subgroup_projects(block["id"], path=block_path)
    except Exception as e:
        logger.error(f"Ошибка при получении блоков: {e}")
    return project_list


def get_project_desc_from_git(session=None, project_id=""):
    project_desc = ""
    try:
        get_custom_attributes_from_git = f"https://gitlab.alfastrah.ru/api/v4/projects/{project_id}/custom_attributes/project_description"
        response = session.get(get_custom_attributes_from_git)
        if response.status_code == 200:
            project_desc = response.json()["value"]
        elif response.status_code == 404:
            logger.info("Атрибут project_description не найден в gitlab.")
    except Exception:
        logger.error("Ошибка: не удалось получить описание проекта из gitlab.")
    return project_desc


def celery_send_omnitracker(sender, title, text, attach, uuid=None):
    bodyNotEncoded = f"""<?xml version="1.0" encoding="utf-8"?>
                <soap:Envelope xmlns:soap="http://schemas.xmlsoap.org/soap/envelope/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance" xmlns:s="http://www.w3.org/2001/XMLSchema" xmlns:tns="http://www.omninet.de/OtWebSvc/v1">
        <soap:Body>
            <tns:InvokeScript>
               <tns:Script name="OmniAPI" runAt="Server">
                    <tns:Parameters>
                        <tns:StringVal name='strFunction'>CreateIncidentFromTemplateWithAttachmentAndExternalNumber</tns:StringVal>
                        <tns:StringVal name='strEMail'>{sender}</tns:StringVal>
                        <tns:StringVal name='strTitle'>{title}</tns:StringVal>
                        <tns:StringVal name='strDescription'>{text}</tns:StringVal>
                        <tns:StringVal name='strTemplateCode'>SH0458</tns:StringVal>
                        <tns:StringVal name='strExternalNumber'>{str(uuid)}</tns:StringVal>
                        <tns:StringVal name='strAttachmentPath'>{attach}</tns:StringVal>
                    </tns:Parameters>
                </tns:Script>
            </tns:InvokeScript>
        </soap:Body>
        </soap:Envelope>
        """

    body = bodyNotEncoded.encode("utf-8")

    try:
        auth = HttpNtlmAuth(
            f"{settings.LDAP_AUTH_CONNECTION_DOMAIN}\\{settings.LDAP_AUTH_CONNECTION_USERNAME}",
            settings.LDAP_AUTH_CONNECTION_PASSWORD,
        )
    except Exception as e:
        print(f"Ошибка авторизации LDAP: {e}")
    result = None

    for i in range(settings.ATTEMPS):
        try:
            if settings.DEBUG:
                res = mock_send_omnitracker_200(
                    taskid="888888",
                    omnitracker_url=settings.OMNITRACKER_URL,
                    body=body,
                    headers=OMNI_HTTP_HEADERS,
                    timeout=settings.TIMEOUT,
                    auth=auth,
                )
            else:
                res = requests.post(
                    settings.OMNITRACKER_URL,
                    body,
                    headers=OMNI_HTTP_HEADERS,
                    timeout=settings.TIMEOUT,
                    auth=auth,
                )

            if res and res.status_code == 200:
                result = res.text
                if result and len(result) > 3 and "strResponse" in result:
                    result = result.split('strResponse">')
                    if len(result) > 1:
                        result = result[1].split("</")[0]
                    else:
                        result = res.text.split("strResponse>")
                        if len(result) > 1:
                            result = result[1].split("</")[0]
                        else:
                            send_to_mattermost(
                                f"[TRACE] Сервер omni вернул неправильный ответ: {res.text}"
                                f"Попытка номер: {i}",
                                f"Ссылка на ACL :https://acl.vesta.ru/acl/info/{str(uuid)}",
                            )
                            result = None
                return result
            else:
                if res.status_code == 401 and not auth:
                    try:
                        auth = HttpNtlmAuth(
                            f"{settings.LDAP_AUTH_CONNECTION_DOMAIN}\\{settings.LDAP_AUTH_CONNECTION_USERNAME}",
                            settings.LDAP_AUTH_CONNECTION_PASSWORD,
                        )
                    except Exception as e:
                        send_to_mattermost(f"Ошибка авторизации LDAP: {e}")
                if i >= settings.ATTEMPS:
                    break
                time.sleep(i)

        except requests.Timeout:
            send_to_mattermost(
                "Подключение не удалось по таймауту (Timeout).",
                f"Попытка подключения номер: {i}.",
                f"OMNITRACKER_URL: {settings.OMNITRACKER_URL}",
                f"Ссылка на ACL: https://acl.vesta.ru/acl/info/{str(uuid)}",
            )
        except requests.ConnectionError:
            send_to_mattermost(
                "Сервис не отвечает (Connection Error).",
                f"Попытка подключения номер: {i}.",
                f"OMNITRACKER_URL: {settings.OMNITRACKER_URL}",
                f"Ссылка на ACL: https://acl.vesta.ru/acl/info/{str(uuid)}",
            )
        except Exception as e:
            send_to_mattermost(
                f"Ошибка Omni request: {e}.",
                f"Попытка номер: {i}.",
                f"OMNITRACKER_URL: {settings.OMNITRACKER_URL}",
                f"Ссылка на ACL: https://acl.vesta.ru/acl/info/ {str(uuid)}",
            )
    return result


def check_acl_in_omni(acl_object):
    """Проверяет acl в omnitracker на наличие заявки"""

    acl_taskid = acl_object.taskid
    if acl_taskid:
        acl_status = acl_object.status
        omni_acl_status, omni_http_status = omni_check_status(acl_taskid)
        if omni_acl_status is not None:
            if acl_status != omni_acl_status:
                acl_object.status = omni_acl_status
                acl_object.save(update_fields=["status"])
            return omni_acl_status
